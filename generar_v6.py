from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

BASE = Path(__file__).resolve().parent
SOURCE = BASE / "ensayo_v6.txt"
OUTPUT = BASE / "Ensayo_organizacion_empresarial_agro_final.docx"

TITLE = "La importancia de la organización empresarial para la agricultura latinoamericana"
HEADINGS = {
    "Comerciantes y empresa mercantil",
    "Innovar de forma organizada",
    "Diferenciarse en el agro",
    "Conclusiones",
}

REFERENCES = [
    [("Brenes, E. R., Montoya, D., & Ciravegna, L. (2014). Differentiation strategies in "
      "emerging markets: The case of Latin American agribusinesses. ", False),
     ("Journal of Business Research, 67", True),
     (", 847-855. https://doi.org/10.1016/j.jbusres.2013.07.003", False)],
    [("Gibson, R., Brenes, E. R., & Barahona, J. C. (2011). Campeones de la innovación en "
      "Latinoamérica. ", False),
     ("INCAE Business Review, 2", True),
     ("(3), 2-6.", False)],
    [("Soto Díaz, R. A. (2025). ", False),
     ("Acuerdos comerciales: Unidad I", True),
     (" [Presentación de PowerPoint]. Escuela Agrícola Panamericana, Zamorano.", False)],
]


def style_run(run, bold=None, italic=None):
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), "Times New Roman")
    lang = OxmlElement("w:lang")
    lang.set(qn("w:val"), "es-ES")
    rpr.append(lang)


def fmt(paragraph, indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    paragraph.alignment = align
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Inches(0.5) if indent else None


def page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    style_run(run)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    val = OxmlElement("w:t")
    val.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, sep, val, end])


def centered_bold(doc, text):
    p = doc.add_paragraph()
    fmt(p, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.keep_with_next = True
    style_run(p.add_run(text), bold=True)


def section_heading(doc, text):
    p = doc.add_paragraph()
    fmt(p, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.keep_with_next = True
    style_run(p.add_run(text), bold=True)


def body(doc, text):
    p = doc.add_paragraph()
    fmt(p)
    style_run(p.add_run(text))


def reference(doc, parts):
    p = doc.add_paragraph()
    fmt(p, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    for text, italic in parts:
        style_run(p.add_run(text), italic=italic)


def main():
    blocks = [b.strip() for b in SOURCE.read_text(encoding="utf-8").split("\n\n") if b.strip()]
    if blocks[0] != TITLE:
        raise SystemExit(f"Título inesperado: {blocks[0]!r}")

    doc = Document()
    sec = doc.sections[0]
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, attr, Inches(1))

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    doc.core_properties.title = TITLE
    doc.core_properties.author = "[Nombre del estudiante]"

    page_number(sec.header.paragraphs[0])

    for _ in range(3):
        fmt(doc.add_paragraph(), indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    centered_bold(doc, TITLE)
    for line in ["[Nombre del estudiante]",
                 "Escuela Agrícola Panamericana, Zamorano",
                 "Acuerdos comerciales",
                 "Dr. Raúl A. Soto D.",
                 "[Fecha de entrega]"]:
        p = doc.add_paragraph()
        fmt(p, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
        style_run(p.add_run(line))
    doc.add_page_break()

    centered_bold(doc, TITLE)

    seen = set()
    for block in blocks[1:]:
        if block == "Referencias":
            doc.add_page_break()
            centered_bold(doc, "Referencias")
            break
        if block in HEADINGS:
            section_heading(doc, block)
            seen.add(block)
        else:
            body(doc, block)

    if HEADINGS - seen:
        raise SystemExit(f"Faltan secciones: {HEADINGS - seen}")

    for parts in REFERENCES:
        reference(doc, parts)

    upd = OxmlElement("w:updateFields")
    upd.set(qn("w:val"), "true")
    doc.settings._element.append(upd)

    doc.save(OUTPUT)
    print(OUTPUT.name)


if __name__ == "__main__":
    main()
