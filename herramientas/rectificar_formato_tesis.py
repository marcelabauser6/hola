#!/usr/bin/env python3
"""Rectificación editorial y visual APA 7 de la tesis auditada.

Etapa 1 (por defecto): normaliza estilos, completa resumen/abstract, reorganiza
las tablas anuales, depura saltos y crea índices estáticos con marcadores.
Etapa 2 (--actualizar-indices PDF): toma las páginas del PDF renderizado y
actualiza los tres índices sin modificar su extensión.
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]
DEFAULT_DOCX = ROOT / "Tesis jaime fredy horacio avance 16 - VERSION AUDITADA.docx"

SUMMARY = (
    "El estudio evaluó la prefactibilidad de producir y distribuir leche entera pasteurizada para un "
    "nicho de mercado en Honduras. Se aplicó un enfoque mixto, organizado con criterios de evaluación "
    "de proyectos de la ONUDI, que integró análisis sectorial y de mercado, estudios técnico, "
    "organizacional, legal y ambiental, una encuesta a 63 establecimientos y una evaluación financiera "
    "a diez años. La propuesta considera una planta en Comayagua, una presentación inicial de un litro "
    "y una comercialización gradual que inicia con 60 clientes activos, 108,000 litros anuales y un "
    "precio de L 28 por litro. La inversión inicial estimada asciende a L 1,784,058 y se financia en "
    "partes iguales con recursos propios y un préstamo. En el escenario determinístico sin "
    "financiamiento, el proyecto obtiene un valor actual neto de L 15,370,631, una tasa interna de "
    "retorno de 55.95 %, un período de recuperación de 3.67 años y una relación beneficio/costo de "
    "1.24. Con financiamiento, el valor actual neto es de L 15,459,686, la tasa interna de retorno "
    "alcanza 63.62 % y la inversión de los socios se recupera en 3.66 años. Los resultados indican que "
    "la propuesta es prefactible bajo los supuestos adoptados. Antes de invertir, deben documentarse "
    "los L 25,000 aún no desagregados en los gastos preoperativos, validar los parámetros probabilísticos "
    "con una nueva corrida de @RISK y asegurar el abastecimiento, la cadena de frío y el cumplimiento "
    "sanitario."
)
KEYWORDS = "Palabras clave: leche pasteurizada, prefactibilidad, agronegocios, evaluación financiera, gestión de riesgos."

ABSTRACT = (
    "This study assessed the prefeasibility of producing and distributing pasteurized whole milk for "
    "a niche market in Honduras. A mixed-method approach, structured around UNIDO project appraisal "
    "criteria, integrated sector and market analyses, technical, organizational, legal, and environmental "
    "studies, a survey of 63 retail establishments, and a ten-year financial evaluation. The proposal "
    "considers a processing plant in Comayagua, an initial one-liter package, and gradual market entry "
    "starting with 60 active customers, annual sales of 108,000 liters, and a price of HNL 28 per liter. "
    "The estimated initial investment is HNL 1,784,058, financed equally with owners’ equity and bank "
    "debt. Under the deterministic unlevered scenario, the project yields a net present value of HNL "
    "15,370,631, an internal rate of return of 55.95%, a 3.67-year payback period, and a benefit-cost "
    "ratio of 1.24. With financing, net present value reaches HNL 15,459,686, internal rate of return "
    "increases to 63.62%, and shareholders recover their investment in 3.66 years. The findings support "
    "the project’s prefeasibility under the assumptions adopted. Before investment, the HNL 25,000 not "
    "yet itemized in preoperating expenses should be documented, the probabilistic parameters should be "
    "validated through a new @RISK run, and raw-milk supply, cold-chain control, and sanitary compliance "
    "should be secured."
)
KEYWORDS_EN = "Keywords: pasteurized milk, prefeasibility, agribusiness, financial appraisal, risk management."

BLACK = RGBColor(0, 0, 0)
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"


def _remove(element) -> None:
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _set_run_font(run, size: float | None = None, italic: bool | None = None, bold: bool | None = None) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:cs"), "Calibri")
    run.font.color.rgb = BLACK
    if size is not None:
        run.font.size = Pt(size)
    if italic is not None:
        run.italic = italic
    if bold is not None:
        run.bold = bold


def _format_all_xml_runs(container, size: float | None = None) -> None:
    for run in container.findall(".//" + qn("w:r")):
        rpr = run.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            run.insert(0, rpr)
        fonts = rpr.find(qn("w:rFonts"))
        if fonts is None:
            fonts = OxmlElement("w:rFonts")
            rpr.insert(0, fonts)
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            fonts.set(qn(f"w:{attr}"), "Calibri")
        for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
            fonts.attrib.pop(qn(f"w:{attr}"), None)
        color = rpr.find(qn("w:color"))
        if color is None:
            color = OxmlElement("w:color")
            rpr.append(color)
        color.set(qn("w:val"), "000000")
        for attr in ("themeColor", "themeTint", "themeShade"):
            color.attrib.pop(qn(f"w:{attr}"), None)
        # Elimina resaltados o sombreados heredados del borrador; el texto
        # editorial debe conservar únicamente negro sobre fondo blanco.
        for tag in ("highlight", "shd"):
            node = rpr.find(qn(f"w:{tag}"))
            if node is not None:
                rpr.remove(node)
        if size is not None:
            for name in ("sz", "szCs"):
                node = rpr.find(qn(f"w:{name}"))
                if node is None:
                    node = OxmlElement(f"w:{name}")
                    rpr.append(node)
                node.set(qn("w:val"), str(int(size * 2)))

    # También normaliza propiedades de carácter heredadas desde el párrafo
    # (w:pPr/w:rPr), que Word y LibreOffice aplican aunque el run no las repita.
    for rpr in container.findall(".//" + qn("w:rPr")):
        fonts = rpr.find(qn("w:rFonts"))
        if fonts is None:
            fonts = OxmlElement("w:rFonts")
            rpr.insert(0, fonts)
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            fonts.set(qn(f"w:{attr}"), "Calibri")
        for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
            fonts.attrib.pop(qn(f"w:{attr}"), None)
        color = rpr.find(qn("w:color"))
        if color is None:
            color = OxmlElement("w:color")
            rpr.append(color)
        color.set(qn("w:val"), "000000")
        for attr in ("themeColor", "themeTint", "themeShade"):
            color.attrib.pop(qn(f"w:{attr}"), None)
        for tag in ("highlight", "shd"):
            node = rpr.find(qn(f"w:{tag}"))
            if node is not None:
                rpr.remove(node)
        if size is not None:
            for name in ("sz", "szCs"):
                node = rpr.find(qn(f"w:{name}"))
                if node is None:
                    node = OxmlElement(f"w:{name}")
                    rpr.append(node)
                node.set(qn("w:val"), str(int(size * 2)))


def _set_style_font(style, size: float, bold: bool = False, italic: bool = False) -> None:
    style.font.name = "Calibri"
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    style.font.color.rgb = BLACK


def _get_or_add_style(doc: Document, name: str, style_type=WD_STYLE_TYPE.PARAGRAPH):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, style_type)


def _configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    _set_style_font(normal, 11)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Inches(0.5)
    pf.widow_control = True

    for level in range(1, 6):
        style = doc.styles[f"Heading {level}"]
        _set_style_font(style, 11, bold=True, italic=level in (3, 5))
        pf = style.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        pf.space_before = Pt(12)
        pf.space_after = Pt(0)
        pf.first_line_indent = None
        pf.keep_with_next = True
        pf.keep_together = True
        pf.page_break_before = level == 1
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
        if level == 4:
            pf.first_line_indent = Inches(0.5)

    caption_number = _get_or_add_style(doc, "APA número de tabla o figura")
    _set_style_font(caption_number, 11, bold=True)
    caption_number.paragraph_format.first_line_indent = None
    caption_number.paragraph_format.line_spacing = 1
    caption_number.paragraph_format.space_before = Pt(8)
    caption_number.paragraph_format.space_after = Pt(0)
    caption_number.paragraph_format.keep_with_next = True

    caption_title = _get_or_add_style(doc, "APA título de tabla o figura")
    _set_style_font(caption_title, 11, italic=True)
    caption_title.paragraph_format.first_line_indent = None
    caption_title.paragraph_format.line_spacing = 1
    caption_title.paragraph_format.space_after = Pt(5)
    caption_title.paragraph_format.keep_with_next = True

    note = _get_or_add_style(doc, "APA nota")
    _set_style_font(note, 10)
    note.paragraph_format.first_line_indent = None
    note.paragraph_format.line_spacing = 1
    note.paragraph_format.space_before = Pt(3)
    note.paragraph_format.space_after = Pt(8)

    continuation = _get_or_add_style(doc, "APA continuación de tabla")
    _set_style_font(continuation, 10, bold=True)
    continuation.paragraph_format.first_line_indent = None
    continuation.paragraph_format.line_spacing = 1
    continuation.paragraph_format.space_before = Pt(8)
    continuation.paragraph_format.space_after = Pt(4)
    continuation.paragraph_format.keep_with_next = True

    for level in (1, 2, 3):
        style = _get_or_add_style(doc, f"Índice manual nivel {level}")
        _set_style_font(style, 10.5)
        pf = style.paragraph_format
        pf.first_line_indent = None
        pf.left_indent = Inches(0.25 * (level - 1))
        pf.line_spacing = 1
        pf.space_after = Pt(2)
        tabs = pf.tab_stops
        tabs.add_tab_stop(Inches(6.45), alignment=2, leader=1)
    for name in ("Índice manual tabla", "Índice manual figura"):
        style = _get_or_add_style(doc, name)
        _set_style_font(style, 10)
        pf = style.paragraph_format
        pf.first_line_indent = None
        pf.line_spacing = 1
        pf.space_after = Pt(2)
        pf.tab_stops.add_tab_stop(Inches(6.45), alignment=2, leader=1)


def _paragraph_after(reference: Paragraph, text: str = "", style=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    reference._p.addnext(new_p)
    paragraph = Paragraph(new_p, reference._parent)
    if style is not None:
        paragraph.style = style
    if text:
        paragraph.add_run(text)
    return paragraph


def _paragraph_before_element(doc: Document, element, text: str, style) -> Paragraph:
    new_p = OxmlElement("w:p")
    element.addprevious(new_p)
    paragraph = Paragraph(new_p, doc._body)
    paragraph.style = style
    paragraph.add_run(text)
    return paragraph


def _find_paragraph(doc: Document, exact: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == exact:
            return paragraph
    raise KeyError(exact)


def _replace_text(paragraph: Paragraph, text: str) -> None:
    for run in list(paragraph.runs):
        _remove(run._element)
    paragraph.add_run(text)


def _remove_empty_paragraphs(doc: Document) -> int:
    removed = 0
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip():
            continue
        p = paragraph._p
        if p.find(".//" + qn("w:sectPr")) is not None:
            continue
        if p.find(".//" + qn("w:drawing")) is not None or p.find(".//" + qn("w:pict")) is not None:
            continue
        _remove(p)
        removed += 1
    return removed


def _complete_summary_abstract(doc: Document) -> None:
    summary = _find_paragraph(doc, "Resumen")
    abstract = _find_paragraph(doc, "Abstract")
    introduction = _find_paragraph(doc, "Introducción")
    for paragraph in list(doc.paragraphs):
        if paragraph._p.getparent() is not summary._p.getparent():
            continue
        children = list(doc.element.body)
        try:
            pos = children.index(paragraph._p)
            s_pos = children.index(summary._p)
            a_pos = children.index(abstract._p)
            i_pos = children.index(introduction._p)
        except ValueError:
            continue
        if s_pos < pos < a_pos or a_pos < pos < i_pos:
            _remove(paragraph._p)
    p1 = _paragraph_after(summary, SUMMARY)
    p2 = _paragraph_after(p1, KEYWORDS)
    p2.paragraph_format.first_line_indent = None
    p2.runs[0].italic = True
    a1 = _paragraph_after(abstract, ABSTRACT)
    a2 = _paragraph_after(a1, KEYWORDS_EN)
    a2.paragraph_format.first_line_indent = None
    a2.runs[0].italic = True


def _assign_headings(doc: Document) -> None:
    level1 = {
        "Resumen", "Abstract", "Introducción", "Objetivos", "Justificación del Proyecto", "Marco Teórico",
        "Metodología", "DESARROLLO DEL ESTUDIO", "Estudio sectorial", "Estudio de Mercado", "Estudio Técnico",
        "Estudio Organizacional y Administrativo", "Estudio Legal", "Estudio ambiental", "Estudio Financiero",
        "Análisis de Riesgos", "Bibliografía",
    }
    level2 = {
        "Plan de Inversiones y Financiamiento", "Producción y Ventas", "Costos Variables", "Costo de Personal.",
        "Costos Fijos.", "Depreciaciones y Amortizaciones.", "Valor de Desecho y Capital de Trabajo.",
        "Tabla de Amortización del Préstamo.", "Impuesto Municipal de Industria, Comercio y Servicios (ICS)",
        "Flujo de Caja del Proyecto (Sin Financiamiento)", "Flujo de Caja del Proyecto con Financiamiento",
        "Costo de Capital sin Financiamiento.", "Indicadores Financieros Sin Financiamiento.",
        "Costo de Capital con Financiamiento (WACC)", "Indicadores Financieros con Financiamiento.",
        "Notas Legales y Cumplimiento Normativo.", "Parametrización para el análisis de sensibilidad y la simulación",
        "Identificación y parametrización de los riesgos", "Análisis de la frecuencia de ocurrencia de los eventos de riesgo",
        "Análisis de la severidad: valor presente de las pérdidas por evento",
        "Cuantificación integrada del riesgo: valor esperado y valor en riesgo",
        "Efecto del riesgo sobre el valor del proyecto", "Limitaciones del modelo de riesgo",
        "Fundamentos de la evaluación de proyectos", "Indicadores económicos clave que afectan el consumo",
        "Resumen del estudio técnico", "Riesgos asociados con la estrategia de nicho",
    }
    level3 = {
        "Análisis de los indicadores financieros sin financiamiento.", "Análisis de los indicadores financieros",
        "Alcance de la parametrización", "Variables Sujetas a Distribución de Probabilidad (@RISK).",
        "Justificación de las distribuciones utilizadas.", "Criterio de lectura de la simulación Monte Carlo",
    }
    level4 = {
        "Valor Actual Neto (VAN)", "Tasa Interna de Retorno (TIR)", "Período de Recuperación de la Inversión (PRI)",
        "Índice de Rentabilidad sobre la Inversión (IRVA)", "Relación Beneficio/Costo", "Inflación anual",
        "Costo de leche cruda", "Merma de proceso", "Costos de empaque", "Costos de combustible",
        "Costos fijos operativos", "Consumo promedio por cliente", "Impuesto sobre la Renta (ISR).",
        "Depreciación y amortización.", "Seguridad social y aportes patronales.", "Prestaciones laborales.",
        "Impuesto sobre Ventas (ISV)", "Arrastre de pérdidas fiscales.",
        "Impuesto Municipal de Industria, Comercio y Servicios", "Capital de trabajo.",
    }
    for p in doc.paragraphs:
        text = p.text.strip()
        if text in level1:
            p.style = doc.styles["Heading 1"]
        elif text in level2:
            p.style = doc.styles["Heading 2"]
        elif text in level3:
            p.style = doc.styles["Heading 3"]
        elif text in level4:
            p.style = doc.styles["Heading 4"]

    # Un mismo nombre aparece en el marco teórico, la metodología y los resultados.
    # Las variantes de mayúsculas no deben alterar la jerarquía: antes del desarrollo
    # son apartados de nivel 2; dentro de los resultados son capítulos de nivel 1.
    chapter_names = {
        "estudio sectorial", "estudio de mercado", "estudio técnico",
        "estudio organizacional y administrativo", "estudio legal",
        "estudio ambiental", "estudio financiero", "análisis de riesgos",
    }
    development_index = next(
        i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "DESARROLLO DEL ESTUDIO"
    )
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().casefold() in chapter_names:
            p.style = doc.styles["Heading 1" if i > development_index else "Heading 2"]


def _revise_sensitivity_and_risk(doc: Document) -> None:
    replacements = {
        "Riesgos  Colocarlo en la parte de riesgos": "Riesgos asociados con la estrategia de nicho",
        "Riesgos Colocarlo en la parte de riesgos": "Riesgos asociados con la estrategia de nicho",
        "Clasificación de Variables para el Análisis de Sensibilidad.": "Parametrización para el análisis de sensibilidad y la simulación",
        "El análisis de sensibilidad tiene como finalidad evaluar el impacto que pueden generar las variaciones de determinadas variables sobre los resultados financieros del proyecto. Debido a que algunos parámetros están sujetos a incertidumbre y cambios en las condiciones del mercado, resulta necesario identificar cuáles variables pueden afectar significativamente la rentabilidad de la inversión.":
            "Esta sección documenta las variables que el modelo trata como fijas y las entradas a las que asigna incertidumbre. La clasificación permite preparar escenarios y simulaciones, pero no demuestra por sí sola qué variable ejerce mayor influencia sobre el VAN o la TIR.",
        "Para este análisis, los supuestos fueron clasificados en variables fijas y variables sensibles. Las variables fijas corresponden a parámetros legales, regulatorios o decisiones internas del proyecto que permanecen constantes durante el horizonte de evaluación. Por su parte, las variables sensibles son aquellas que pueden experimentar cambios debido a factores de mercado, condiciones económicas o desempeño operativo.":
            "Las variables fijas corresponden a parámetros normativos, contractuales o decisiones internas que no se simulan. Las variables probabilísticas representan condiciones de mercado, operación y demanda mediante los rangos y distribuciones configurados en el libro financiero. Los valores se presentan como mínimo, base y máximo para que su origen sea verificable.",
        "Variables más críticas del proyecto": "Alcance de la parametrización",
        "Variables de mayor impacto potencial": "La tabla identifica entradas inciertas; no constituye un ranking de sensibilidad.",
        "De acuerdo con la naturaleza del proyecto, las variables que podrían generar un mayor impacto sobre la rentabilidad son:":
            "Para establecer una jerarquía de impacto se requiere una salida específica, como un gráfico tornado, coeficientes de correlación de Spearman o una regresión estandarizada. Esa salida no está incorporada en el expediente; por ello, no se atribuye un orden de importancia cuantitativo a las variables.",
        "Costo de leche cruda, por representar el principal componente del costo de producción.":
            "El costo de la leche, el consumo, el precio, el número de clientes, la merma, el empaque, el combustible, la inflación y los costos fijos se mantienen como entradas relevantes por su relación directa con ingresos o egresos, sin afirmar cuál es la más influyente.",
        "Consumo promedio por cliente, debido a su efecto directo sobre el nivel de ventas e ingresos.": "",
        "Costo de combustible de reparto, por su alta volatilidad en el mercado.": "",
        "Merma de proceso, al influir directamente en la eficiencia productiva.": "",
        "Inflación anual, por afectar simultáneamente los costos operativos y los precios de venta.": "",
        "Estas variables fueron seleccionadas para realizar escenarios alternativos y evaluar la estabilidad financiera del proyecto frente a distintos niveles de riesgo.":
            "La interpretación de sensibilidad deberá completarse cuando se genere una salida reproducible que relacione cada entrada con un indicador financiero de respuesta.",
        "Importancia del análisis Monte Carlo.": "Criterio de lectura de la simulación Monte Carlo",
        "La simulación Monte Carlo permite generar miles de escenarios posibles a partir de las distribuciones de probabilidad asignadas a cada variable. De esta manera, es posible estimar la probabilidad de obtener determinados niveles de rentabilidad, así como identificar los factores que ejercen mayor influencia sobre indicadores financieros como el Valor Actual Neto (VAN) y la Tasa Interna de Retorno (TIR).":
            "La simulación Monte Carlo combina las distribuciones de entrada para obtener una distribución de posibles resultados. Una corrida permite describir rangos, medias y percentiles del VAN o la TIR; para identificar influencia relativa debe acompañarse de una medida de sensibilidad. En esta tesis se conservan los resultados archivados de @RISK, pero no se presenta un tornado ni un coeficiente de sensibilidad, por lo que no se reporta un ranking.",
        "Este enfoque proporciona una visión más robusta del riesgo del proyecto, complementando los resultados obtenidos mediante la evaluación financiera tradicional y facilitando la toma de decisiones bajo condiciones de incertidumbre.":
            "Los resultados probabilísticos deben leerse como complemento del escenario determinístico y validarse mediante una nueva corrida de @RISK antes de la decisión final de inversión.",
    }
    to_delete = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text in replacements:
            new = replacements[text]
            if new:
                _replace_text(p, new)
            else:
                to_delete.append(p._p)
    for element in to_delete:
        _remove(element)

    risk_intro = _find_paragraph(doc, "Es pertinente señalar dos aspectos relativos a la presentación de los resultados. En primer lugar, las gráficas fueron generadas con la versión académica del programa, la cual incorpora una marca de agua identificatoria en cada figura y rotula el eje de valores con el símbolo genérico de moneda; no obstante, todas las cifras reportadas están expresadas en lempiras. En segundo lugar, las bandas delimitadas en la parte superior de cada gráfica corresponden al intervalo de confianza del 90 %, definido por los percentiles 5 y 95 de la distribución simulada, siendo este último el que se emplea posteriormente como valor en riesgo al 95 % de confianza.")
    _replace_text(risk_intro,
        "Las gráficas proceden de la versión académica de @RISK y conservan su marca de agua. En las gráficas de una sola distribución cuya franja superior muestra 5.0 %–90.0 %–5.0 %, el 90.0 % no significa ‘percentil 90’: identifica el intervalo central que contiene el 90 % de las iteraciones. El separador izquierdo es el percentil 5 (P5) y el derecho es el percentil 95 (P95); queda 5.0 % de los resultados en cada cola. En las distribuciones de pérdidas, el P95 se reporta como VaR al 95 %. El VaR al 99 % de la Tabla 35 proviene del resumen numérico de la simulación y no está señalado en esa franja. Los porcentajes del eje vertical representan frecuencia relativa por clase, no niveles de confianza. Todas las cantidades monetarias están expresadas en lempiras.")

    freq = _find_paragraph(doc, "El análisis de la frecuencia resulta indispensable porque determina el carácter con que cada riesgo debe ser gestionado. Los eventos de frecuencia elevada constituyen condiciones habituales de la operación y deben incorporarse a la estructura de costos y a los procedimientos ordinarios de control, mientras que los eventos de frecuencia reducida configuran contingencias excepcionales cuyo tratamiento apropiado suele ser la transferencia del riesgo a un tercero.")
    _replace_text(freq,
        "La lectura de frecuencia combina dos fuentes. El valor esperado de ocurrencias se obtiene sumando las probabilidades anuales de la Tabla 34, porque cada año se modela con una variable Bernoulli que solo admite 0 o 1 evento. La probabilidad de cero ocurrencias se obtiene multiplicando (1 − p) para los diez años. Los límites P5 y P95 se leen en los dos separadores de la franja roja de cada figura. Así, todo porcentaje o promedio citado puede rastrearse a la tabla de parámetros o a la salida gráfica.")

    for p in doc.paragraphs:
        if p.text and "intervalo de confianza del 90 %" in p.text:
            _replace_text(p, p.text.replace("intervalo de confianza del 90 %", "intervalo central del 90 %"))
        elif p.text and "intervalo de confianza del 90%" in p.text:
            _replace_text(p, p.text.replace("intervalo de confianza del 90%", "intervalo central del 90 %"))

    for p in doc.paragraphs:
        if p.text.startswith("Corresponde formular una precisión metodológica relativa a la comparación"):
            _replace_text(p,
                "La media probabilística del VPN del inversionista (L 23,999,077) es mayor que el VAN determinístico con financiamiento (L 15,459,686). El expediente no contiene una corrida reproducible que permita atribuir con certeza esa diferencia a una variable específica. Por ello, la cifra probabilística se presenta como resultado archivado del modelo y no sustituye el VAN determinístico en la decisión. Antes de la presentación definitiva deben recalibrarse las entradas, ejecutar nuevamente @RISK y comprobar que la media simulada sea coherente con el escenario base.")
        if p.text.startswith("Estas limitaciones no comprometen la validez de los resultados presentados"):
            _replace_text(p,
                "Estas limitaciones impiden certificar los resultados probabilísticos como definitivos. Los valores se conservan para mostrar la metodología y el orden de magnitud de la exposición; la decisión de inversión debe apoyarse en los resultados determinísticos hasta que una nueva corrida documentada confirme medias, percentiles y valores en riesgo.")
        if p.text.startswith("El hallazgo central del análisis es la separación completa"):
            _replace_text(p,
                "El hallazgo central es la separación completa de ambas distribuciones. El máximo simulado de VPN Risk (L 11,934,048.84) es inferior al mínimo simulado de VPN Inversionista (L 22,095,576.95). Por tanto, dentro de los rangos archivados de esta corrida, la pérdida agregada no alcanza el valor del proyecto. Esta conclusión procede de los extremos numéricos visibles en la Figura 54 y transcritos en la Tabla 36; no depende de interpretar las etiquetas porcentuales de los separadores.")
        if p.text.startswith("Un segundo elemento digno de mención es la diferencia en la dispersión relativa"):
            _replace_text(p,
                "La Tabla 36, y no la altura visual de las curvas superpuestas, contiene los estadísticos comparables. A partir de sus medias y desviaciones estándar se obtiene un coeficiente de variación de 41.5 % para VPN Risk y de 2.1 % para VPN Inversionista; ambos se calculan como desviación estándar dividida entre media. Las etiquetas 87.6 %, 5.0 %, 0.0 % y 100.0 %, junto con los valores 1.70 y 6.37 visibles en la franja superior, forman parte de la captura superpuesta de @RISK y están asociadas a sus separadores verticales. La imagen estática no identifica de manera inequívoca la serie y el rango a los que corresponde cada etiqueta; por ello, no se interpretan como niveles de confianza ni se utilizan en la comparación. La conclusión se sustenta en los mínimos, máximos, medias y desviaciones estándar visibles en la propia figura y transcritos en la Tabla 36.")


def _matrix(table: Table) -> list[list[str]]:
    return [[cell.text.strip() for cell in row.cells] for row in table.rows]


def _new_table_before(doc: Document, reference, data: list[list[str]]) -> Table:
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    for i, row in enumerate(data):
        for j, value in enumerate(row):
            table.cell(i, j).text = value
    reference.addprevious(table._tbl)
    return table


def _table_number_map(doc: Document) -> dict[int, Table]:
    result: dict[int, Table] = {}
    pending: int | None = None
    for child in doc.element.body:
        if child.tag == qn("w:p"):
            text = Paragraph(child, doc._body).text.strip()
            match = re.fullmatch(r"Tabla (\d+)", text)
            pending = int(match.group(1)) if match else pending
        elif child.tag == qn("w:tbl") and pending is not None:
            result[pending] = Table(child, doc._body)
            pending = None
    return result


def _set_cell_border(cell, edge: str, value: str, size: str = "0") -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    borders = tcpr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcpr.append(borders)
    node = borders.find(qn(f"w:{edge}"))
    if node is None:
        node = OxmlElement(f"w:{edge}")
        borders.append(node)
    node.set(qn("w:val"), value)
    node.set(qn("w:sz"), size)
    node.set(qn("w:color"), "000000")
    node.set(qn("w:space"), "0")


def _format_table(table: Table, widths: list[float] | None = None, size: float = 10) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = widths is None
    tblpr = table._tbl.tblPr
    layout = tblpr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblpr.append(layout)
    layout.set(qn("w:type"), "autofit" if widths is None else "fixed")
    borders = tblpr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblpr.append(borders)
    for edge, val, sz in (("top", "single", "8"), ("bottom", "single", "8"), ("left", "nil", "0"), ("right", "nil", "0"), ("insideH", "nil", "0"), ("insideV", "nil", "0")):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), val)
        node.set(qn("w:sz"), sz)
        node.set(qn("w:color"), "000000")
    for ri, row in enumerate(table.rows):
        trpr = row._tr.get_or_add_trPr()
        for old in trpr.findall(qn("w:cantSplit")):
            trpr.remove(old)
        if ri == 0:
            header = trpr.find(qn("w:tblHeader"))
            if header is None:
                header = OxmlElement("w:tblHeader")
                trpr.append(header)
        for ci, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if widths:
                cell.width = Inches(widths[ci])
            tcpr = cell._tc.get_or_add_tcPr()
            shd = tcpr.find(qn("w:shd"))
            if shd is not None:
                tcpr.remove(shd)
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                _set_cell_border(cell, edge, "nil")
            if ri == 0:
                _set_cell_border(cell, "bottom", "single", "8")
            margins = tcpr.find(qn("w:tcMar"))
            if margins is None:
                margins = OxmlElement("w:tcMar")
                tcpr.append(margins)
            for edge in ("top", "start", "bottom", "end"):
                node = margins.find(qn(f"w:{edge}"))
                if node is None:
                    node = OxmlElement(f"w:{edge}")
                    margins.append(node)
                node.set(qn("w:w"), "55")
                node.set(qn("w:type"), "dxa")
            for p in cell.paragraphs:
                p.paragraph_format.first_line_indent = None
                p.paragraph_format.line_spacing = 1
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else (WD_ALIGN_PARAGRAPH.CENTER if ri == 0 else WD_ALIGN_PARAGRAPH.RIGHT)
                for run in p.runs:
                    _set_run_font(run, size=size, bold=True if ri == 0 else None)


def _replace_table_with_column_blocks(
    doc: Document,
    table: Table,
    number: int,
    blocks: list[tuple[list[int], str | None, list[float]]],
    size: float = 10,
) -> None:
    """Reemplaza una tabla densa por bloques legibles que repiten descriptores."""
    data = _matrix(table)
    for columns, continuation, widths in blocks:
        if continuation:
            _paragraph_before_element(
                doc,
                table._tbl,
                f"Tabla {number} (continuación). {continuation}",
                doc.styles["APA continuación de tabla"],
            )
        block = [[row[index] for index in columns] for row in data]
        new_table = _new_table_before(doc, table._tbl, block)
        _format_table(new_table, widths, size)
    _remove(table._tbl)


def _split_wide_tables(doc: Document) -> None:
    mapping = _table_number_map(doc)

    # Tabla 13: diez años sin columna de Año 0.
    _replace_table_with_column_blocks(
        doc,
        mapping[13],
        13,
        [
            ([0, 1, 2, 3, 4], None, [2.05, 1.075, 1.075, 1.075, 1.075]),
            ([0, 5, 6, 7], "Años 5–7", [2.05, 1.43, 1.43, 1.43]),
            ([0, 8, 9, 10], "Años 8–10", [2.05, 1.43, 1.43, 1.43]),
        ],
    )

    # Tablas financieras con Año 0–10. Cuatro o tres años por bloque evitan
    # que cifras de siete u ocho dígitos se partan entre líneas.
    annual_tables = [16, 17, 19, 20, 22, 23, 24, 25, 26]
    for number in annual_tables:
        _replace_table_with_column_blocks(
            doc,
            mapping[number],
            number,
            [
                ([0, 1, 2, 3, 4], None, [2.05, 1.075, 1.075, 1.075, 1.075]),
                ([0, 5, 6, 7], "Años 4–6", [2.05, 1.43, 1.43, 1.43]),
                ([0, 8, 9, 10, 11], "Años 7–10", [2.05, 1.075, 1.075, 1.075, 1.075]),
            ],
        )

    # La Tabla 21 conserva Activo y Valor antes de los años proyectados.
    _replace_table_with_column_blocks(
        doc,
        mapping[21],
        21,
        [
            ([0, 1, 2, 3, 4, 5], None, [1.75, 0.95, 0.91, 0.91, 0.91, 0.91]),
            ([0, 1, 6, 7, 8], "Años 5–7", [1.75, 0.95, 1.22, 1.22, 1.22]),
            ([0, 1, 9, 10, 11], "Años 8–10", [1.75, 0.95, 1.22, 1.22, 1.22]),
        ],
    )


def _split_dense_tables(doc: Document) -> None:
    """Separa columnas textuales o estadísticas que resultan ilegibles juntas."""
    mapping = _table_number_map(doc)
    _replace_table_with_column_blocks(
        doc,
        mapping[5],
        5,
        [
            ([0, 1, 2], None, [1.55, 0.85, 4.10]),
            ([0, 3, 4], "Justificación y estrategia", [1.30, 2.60, 2.60]),
        ],
    )
    _replace_table_with_column_blocks(
        doc,
        mapping[15],
        15,
        [
            ([0, 1, 4, 5], None, [2.80, 1.35, 1.05, 1.30]),
            ([0, 2, 3], "Estructura de financiamiento", [3.20, 1.65, 1.65]),
        ],
    )
    _replace_table_with_column_blocks(
        doc,
        mapping[34],
        34,
        [
            ([0, 1, 2], None, [2.75, 1.90, 1.85]),
            ([0, 3, 4, 5], "Parámetros de impacto", [2.75, 1.25, 1.25, 1.25]),
        ],
    )
    _replace_table_with_column_blocks(
        doc,
        mapping[35],
        35,
        [
            ([0, 1, 4], None, [3.35, 1.65, 1.50]),
            ([0, 2, 3, 5], "Valores en riesgo", [2.75, 1.25, 1.25, 1.25]),
        ],
    )


def _rebuild_sensitivity_tables(doc: Document) -> None:
    mapping = _table_number_map(doc)
    t32 = mapping[32]
    source32 = _matrix(t32)
    data32 = [["Parámetro", "Valor base", "Clasificación", "Rango o criterio aplicado"]]
    for row in source32[1:]:
        criterion = "No se simula. " + row[6] if row[2].lower().startswith("fija") else f"{row[3]} – {row[4]} – {row[5]}. {row[6]}"
        data32.append([row[0], row[1], row[2], criterion])
    new32 = _new_table_before(doc, t32._tbl, data32)
    _remove(t32._tbl)
    _format_table(new32, [2.05, 1.00, 1.00, 2.45], 10)

    mapping = _table_number_map(doc)
    t33 = mapping[33]
    source33 = _matrix(t33)
    data33 = [["Parámetro y unidad", "Base", "Mínimo", "Máximo", "Distribución configurada"]]
    for row in source33[1:]:
        label = f"{row[0]} ({row[5]})" if row[5] and row[5] != "%" else (f"{row[0]} (%)" if row[5] == "%" else row[0])
        data33.append([label, row[1], row[2], row[3], row[4]])
    new33 = _new_table_before(doc, t33._tbl, data33)
    _remove(t33._tbl)
    _format_table(new33, [2.15, 0.75, 0.75, 0.75, 2.10], 10)


def _format_tables_and_captions(doc: Document) -> None:
    for table in doc.tables:
        if not any(table is t for t in []):
            cols = len(table.columns)
            size = 9.5 if cols >= 6 else 10
            _format_table(table, None, size)
    for p in doc.paragraphs:
        text = p.text.strip()
        if re.fullmatch(r"(?:Tabla|Figura) \d+", text):
            p.style = doc.styles["APA número de tabla o figura"]
            for run in p.runs:
                _set_run_font(run, 11, bold=True)
            node = p._p.getnext()
            while node is not None and node.tag != qn("w:p"):
                node = node.getnext()
            if node is not None:
                title = Paragraph(node, p._parent)
                title.style = doc.styles["APA título de tabla o figura"]
                for run in title.runs:
                    _set_run_font(run, 11, italic=True)
        elif text.startswith("Nota."):
            p.style = doc.styles["APA nota"]
            for run in p.runs:
                _set_run_font(run, 10)
            if p.runs:
                p.runs[0].italic = True
        elif text.startswith("Tabla ") and "(continuación)" in text:
            p.style = doc.styles["APA continuación de tabla"]


def _resize_figures(doc: Document) -> None:
    max_w = Inches(6.25)
    max_h = Inches(6.6)
    for shape in doc.inline_shapes:
        ratio = min(1.0, max_w / shape.width, max_h / shape.height)
        if ratio < 1:
            new_width = int(shape.width * ratio)
            new_height = int(shape.height * ratio)
            try:
                shape.width = new_width
                shape.height = new_height
            except AttributeError:
                # SmartArt/diagramas no exponen pic.spPr, pero sí el extent inline.
                shape._inline.extent.cx = new_width
                shape._inline.extent.cy = new_height
    for p in doc.paragraphs:
        if p._p.find(".//" + qn("w:drawing")) is not None:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = None
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_together = True


def _add_page_numbers(doc: Document) -> None:
    for i, section in enumerate(doc.sections):
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.5)
        section.footer_distance = Inches(0.5)
        header = section.header
        header.is_linked_to_previous = i > 0
        if i == 0:
            for p in list(header.paragraphs):
                _remove(p._p)
            p = header.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.first_line_indent = None
            field = OxmlElement("w:fldSimple")
            field.set(qn("w:instr"), "PAGE")
            run = OxmlElement("w:r")
            text = OxmlElement("w:t")
            text.text = "1"
            run.append(text)
            field.append(run)
            p._p.append(field)
            _format_all_xml_runs(p._p, 11)


def _normalize_body_paragraphs(doc: Document) -> None:
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if p.style.name.startswith("Heading") or p.style.name.startswith("APA ") or p.style.name.startswith("Índice manual"):
            continue
        if p._p.find(".//" + qn("w:drawing")) is not None:
            continue
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        if not text.startswith(("•", "-")) and p.style.name != "List Paragraph":
            p.paragraph_format.first_line_indent = Inches(0.5)
        for run in p.runs:
            _set_run_font(run, 11)
    for section in doc.sections:
        _format_all_xml_runs(section.header._element)
        _format_all_xml_runs(section.footer._element)
    _format_all_xml_runs(doc.element.body)


def _caption_entries(doc: Document, kind: str) -> list[tuple[str, str]]:
    result = []
    pattern = re.compile(rf"{kind} (\d+)$")
    for p in doc.paragraphs:
        match = pattern.fullmatch(p.text.strip())
        if not match:
            continue
        node = p._p.getnext()
        while node is not None and node.tag != qn("w:p"):
            node = node.getnext()
        title = Paragraph(node, p._parent).text.strip() if node is not None else ""
        result.append((f"{kind} {match.group(1)}", title))
    return result


def _heading_entries(doc: Document) -> list[tuple[int, str]]:
    entries = []
    for p in doc.paragraphs:
        if p.style.name in {"Heading 1", "Heading 2", "Heading 3"}:
            text = p.text.strip()
            if text and not text.casefold().startswith("índice de") and text != "Tabla de contenido":
                entries.append((int(p.style.name[-1]), text))
    return entries


def _clear_front_indexes(doc: Document) -> None:
    body = doc.element.body
    for child in list(body):
        if child.tag == qn("w:sdt"):
            gallery = child.find(".//" + qn("w:docPartGallery"))
            if gallery is not None and "Table of Contents" in gallery.get(qn("w:val"), ""):
                _remove(child)
    summary = _find_paragraph(doc, "Resumen")._p
    start = None
    for p in doc.paragraphs:
        if p.text.strip() in {"Tabla de contenido", "Índice de Tablas", "Índice de tablas"}:
            start = p._p
            break
    if start is not None:
        children = list(body)
        for child in children[children.index(start):children.index(summary)]:
            _remove(child)


def _insert_front_indexes(doc: Document, headings, tables, figures) -> None:
    summary = _find_paragraph(doc, "Resumen")
    target = summary._p
    blocks: list[Paragraph] = []
    title = _paragraph_before_element(doc, target, "Tabla de contenido", doc.styles["Heading 1"])
    blocks.append(title)
    for level, text in headings:
        blocks.append(_paragraph_before_element(doc, target, f"{text}\t000", doc.styles[f"Índice manual nivel {level}"]))
    title = _paragraph_before_element(doc, target, "Índice de tablas", doc.styles["Heading 1"])
    blocks.append(title)
    for number, text in tables:
        blocks.append(_paragraph_before_element(doc, target, f"{number}. {text}\t000", doc.styles["Índice manual tabla"]))
    title = _paragraph_before_element(doc, target, "Índice de figuras", doc.styles["Heading 1"])
    blocks.append(title)
    for number, text in figures:
        blocks.append(_paragraph_before_element(doc, target, f"{number}. {text}\t000", doc.styles["Índice manual figura"]))
    for p in blocks:
        for run in p.runs:
            _set_run_font(run, 10.5 if "nivel" in p.style.name else 10)


def _settings_update_fields(doc: Document) -> None:
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def rectificar(docx_path: Path) -> None:
    doc = Document(docx_path)
    _configure_styles(doc)
    _complete_summary_abstract(doc)
    _revise_sensitivity_and_risk(doc)
    _remove_empty_paragraphs(doc)
    _assign_headings(doc)
    headings = _heading_entries(doc)
    tables = _caption_entries(doc, "Tabla")
    figures = _caption_entries(doc, "Figura")
    if len(tables) != 37 or len(figures) != 54:
        raise RuntimeError(f"Inventario inesperado: {len(tables)} tablas y {len(figures)} figuras")
    _clear_front_indexes(doc)
    _format_tables_and_captions(doc)
    _split_wide_tables(doc)
    _split_dense_tables(doc)
    _rebuild_sensitivity_tables(doc)
    _resize_figures(doc)
    _add_page_numbers(doc)
    _normalize_body_paragraphs(doc)
    _insert_front_indexes(doc, headings, tables, figures)
    _settings_update_fields(doc)
    doc.core_properties.title = "Estudio de prefactibilidad para la producción y distribución de leche pasteurizada"
    doc.core_properties.subject = "Versión auditada y rectificada con formato APA 7"
    doc.core_properties.comments = "Fuente Calibri, índices actualizados, tablas legibles y resultados probabilísticos sujetos a recalibración documentada."
    doc.save(docx_path)


def _norm(text: str) -> str:
    return " ".join(text.replace("\u00a0", " ").split()).casefold()


def _find_page(page_texts: list[str], needle: str, start: int) -> int:
    target = _norm(needle)
    short = target[:90]
    for i in range(max(0, start), len(page_texts)):
        if target in page_texts[i] or (len(short) > 20 and short in page_texts[i]):
            return i
    for i in range(max(0, start), len(page_texts)):
        words = [w for w in re.split(r"\W+", short) if len(w) > 4]
        if words and sum(w in page_texts[i] for w in words) >= max(2, int(len(words) * 0.7)):
            return i
    raise RuntimeError(f"No se localizó en PDF: {needle}")


def actualizar_indices(docx_path: Path, pdf_path: Path) -> None:
    try:
        import fitz
    except ImportError as exc:
        raise RuntimeError("Se requiere PyMuPDF para actualizar los índices") from exc
    pdf = fitz.open(pdf_path)
    page_texts = [_norm(page.get_text()) for page in pdf]
    doc = Document(docx_path)
    summary_page = _find_page(page_texts, SUMMARY[:100], 0)

    groups = [
        # El encabezado «Resumen» puede quedar en la página anterior a su texto.
        ("Índice manual nivel", max(0, summary_page - 2)),
        ("Índice manual tabla", summary_page),
        ("Índice manual figura", summary_page),
    ]
    for prefix, initial in groups:
        cursor = initial
        paragraphs = [p for p in doc.paragraphs if p.style.name.startswith(prefix)]
        for p in paragraphs:
            label = p.text.rsplit("\t", 1)[0]
            if prefix == "Índice manual tabla":
                search = label.split(".", 1)[0]
            elif prefix == "Índice manual figura":
                search = label.split(".", 1)[0]
            else:
                search = label
            page_index = _find_page(page_texts, search, cursor)
            cursor = page_index
            _replace_text(p, f"{label}\t{page_index + 1}")
            for run in p.runs:
                _set_run_font(run, 10.5 if prefix.endswith("nivel") else 10)
    doc.save(docx_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--docx", type=Path, default=DEFAULT_DOCX)
    parser.add_argument("--actualizar-indices", type=Path, metavar="PDF")
    args = parser.parse_args()
    if args.actualizar_indices:
        actualizar_indices(args.docx, args.actualizar_indices)
        print(f"Índices actualizados: {args.docx.name}")
    else:
        rectificar(args.docx)
        print(f"Formato APA 7 rectificado: {args.docx.name}")


if __name__ == "__main__":
    main()
