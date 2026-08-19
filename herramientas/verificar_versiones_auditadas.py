#!/usr/bin/env python3
"""Verificación estructural y semántica de las versiones auditadas."""

from __future__ import annotations

import hashlib
import math
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

import openpyxl
from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
DOCX_ORIGINAL = ROOT / "Tesis jaime fredy horacio avance 16.docx"
DOCX_AUDITADO = ROOT / "Tesis jaime fredy horacio avance 16 - VERSION AUDITADA.docx"
XLSM_ORIGINAL = ROOT / "Estudio Financiero @risk Jaime Horacio Fredy (1).xlsm"
XLSM_AUDITADO = ROOT / "Estudio Financiero @risk Jaime Horacio Fredy (1) - VERSION AUDITADA.xlsm"

S_NS = "http://schemas.openxmlformats.org/spreadsheetml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
PKG_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"

CHECKS = 0


def check(condition: bool, message: str) -> None:
    global CHECKS
    CHECKS += 1
    if not condition:
        raise AssertionError(message)


def digest(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            h.update(block)
    return h.hexdigest()


def matrix(table) -> list[list[str]]:
    return [[cell.text.strip() for cell in row.cells] for row in table.rows]


def row_by_label(table, label: str) -> list[str]:
    for row in matrix(table):
        if row[0] == label:
            return row
    raise AssertionError(f"No se encontró la fila {label!r}")


def number(text: str) -> float:
    cleaned = text.replace("L", "").replace(",", "").replace("%", "").strip()
    return float(cleaned)


def fills(table) -> list[str]:
    result = []
    for row in table.rows:
        for cell in row.cells:
            shd = cell._tc.get_or_add_tcPr().find(qn("w:shd"))
            if shd is not None:
                value = shd.get(qn("w:fill"))
                if value and value.lower() not in {"auto", "ffffff"}:
                    result.append(value)
    return result


def verify_docx() -> None:
    original = Document(DOCX_ORIGINAL)
    audited = Document(DOCX_AUDITADO)
    check(len(original.tables) == 37, "El original ya no tiene las 37 tablas esperadas")
    check(len(audited.tables) == 37, "La copia auditada no tiene 37 tablas")
    check(len(original.paragraphs) == len(audited.paragraphs) == 2050, "Cambió el número de párrafos")
    check(len(original.sections) == len(audited.sections) == 3, "Cambió el número de secciones")
    for left, right in zip(original.sections, audited.sections):
        check(left.orientation == right.orientation, "Cambió la orientación de una sección")
        check(left.page_width == right.page_width and left.page_height == right.page_height, "Cambió el tamaño de página")

    expected_dimensions = {
        7: (12, 6), 11: (7, 4), 15: (9, 6), 16: (5, 12), 17: (5, 12),
        18: (9, 3), 19: (2, 12), 20: (5, 12), 21: (8, 12), 22: (5, 12),
        23: (6, 12), 24: (8, 12), 25: (17, 12), 26: (20, 12), 27: (3, 5),
        28: (6, 2), 29: (5, 5), 30: (6, 2), 31: (9, 2), 32: (24, 7), 33: (12, 6),
    }
    for table_number, dimensions in expected_dimensions.items():
        table = audited.tables[table_number - 1]
        actual = (len(table.rows), len(table.columns))
        check(actual == dimensions, f"Dimensiones inesperadas en T{table_number}: {actual}")
        check(not fills(table), f"T{table_number} conserva rellenos de color")
        data = matrix(table)
        for column in range(len(data[0])):
            check(any(row[column] for row in data), f"T{table_number} tiene una columna totalmente vacía")
        tbl_borders = table._tbl.tblPr.find(qn("w:tblBorders"))
        check(tbl_borders is not None, f"T{table_number} no tiene definición de bordes")
        check(tbl_borders.find(qn("w:top")).get(qn("w:val")) == "single", f"T{table_number} sin borde superior")
        check(tbl_borders.find(qn("w:bottom")).get(qn("w:val")) == "single", f"T{table_number} sin borde inferior")
        check(tbl_borders.find(qn("w:insideV")).get(qn("w:val")) == "nil", f"T{table_number} conserva líneas verticales")

    t7 = audited.tables[6]
    data7 = matrix(t7)
    for row in data7[1:-1]:
        weight = number(row[1]) / 100
        check(math.isclose(number(row[3]), weight * number(row[2]), abs_tol=0.001), f"Ponderación incorrecta T7: {row[0]}")
        check(math.isclose(number(row[5]), weight * number(row[4]), abs_tol=0.001), f"Ponderación incorrecta T7: {row[0]}")
    check(data7[-1][3] == "2.50" and data7[-1][5] == "1.60", "Totales incorrectos en T7")

    t11 = audited.tables[10]
    check(row_by_label(t11, "Total")[1:] == ["76,000", "", "15,200"], "T11 no quedó conciliada")
    check(any("Provisión no desagregada" in row[0] for row in matrix(t11)), "Falta la provisión transparente en T11")

    check(row_by_label(audited.tables[22], "Cuota total")[7] == "0", "T23 conserva cuota en Año 6")
    check(row_by_label(audited.tables[27], "PRI (años)")[1] == "3.67", "PRI sin financiamiento incorrecto")
    check(row_by_label(audited.tables[29], "PRI (años)")[1] == "3.66", "PRI con financiamiento incorrecto")

    for table_number in (25, 26):
        table = audited.tables[table_number - 1]
        check(any(row[0] == "(−) Impuesto municipal (ICS)" for row in matrix(table)), f"Falta ICS en T{table_number}")
        old_flow = row_by_label(original.tables[table_number - 1], "(=) FLUJO DE CAJA NETO")
        new_flow = row_by_label(table, "(=) FLUJO DE CAJA NETO")
        check(old_flow == new_flow, f"Los flujos netos cambiaron al transparentar ICS en T{table_number}")

    # Reproducibilidad de la utilidad antes de impuestos con cifras publicadas (tolerancia de redondeo: L 2).
    t25 = audited.tables[24]
    revenue = row_by_label(t25, "(+) Ingreso por ventas")
    variable = row_by_label(t25, "(−) Costos variables")
    fixed = row_by_label(t25, "(−) Costos fijos")
    depreciation = row_by_label(t25, "(−) Depreciación y amortización")
    ics = row_by_label(t25, "(−) Impuesto municipal (ICS)")
    uai = row_by_label(t25, "(=) Utilidad antes de impuestos")
    for column in range(1, 12):
        calculated = number(revenue[column]) - number(variable[column]) - number(fixed[column]) - number(depreciation[column]) - number(ics[column])
        check(abs(calculated - number(uai[column])) <= 2, f"T25 no reproduce UAI en columna {column}")

    t26 = audited.tables[25]
    revenue = row_by_label(t26, "(+) Ingreso por ventas")
    variable = row_by_label(t26, "(−) Costos variables")
    fixed = row_by_label(t26, "(−) Costos fijos")
    interest = row_by_label(t26, "(−) Gastos financieros")
    depreciation = row_by_label(t26, "(−) Depreciación y amortización")
    ics = row_by_label(t26, "(−) Impuesto municipal (ICS)")
    uai = row_by_label(t26, "(=) Utilidad antes de impuestos")
    for column in range(1, 12):
        calculated = number(revenue[column]) - number(variable[column]) - number(fixed[column]) - number(interest[column]) - number(depreciation[column]) - number(ics[column])
        check(abs(calculated - number(uai[column])) <= 2, f"T26 no reproduce UAI en columna {column}")

    t32_labels = {row[0] for row in matrix(audited.tables[31])}
    t33_labels = {row[0] for row in matrix(audited.tables[32])}
    for label in ("Precio de venta, Año 1 (L/litro)", "Clientes activos, Año 1", "Clientes activos, Año 5", "Clientes activos, Año 10"):
        check(label in t32_labels, f"T32 omite {label}")
    for label in ("Precio de venta, Año 1", "Clientes activos, Año 1", "Clientes activos, Año 5", "Clientes activos, Año 10"):
        check(label in t33_labels, f"T33 omite {label}")
    check("Normal (media 4.5 %; DE 0.6 %)" in row_by_label(audited.tables[32], "Inflación anual"), "Distribución de inflación incorrecta")

    full_text = "\n".join(p.text for p in audited.paragraphs)
    check("PRI de 4.67" not in full_text and "PRI de 4.66" not in full_text, "Persisten PRI antiguos en la narrativa")
    check("EJEMPLO DADO POR EL ASESOR" not in "\n".join(" | ".join(row) for row in matrix(audited.tables[31]) + matrix(audited.tables[32])), "Persisten notas internas en T32/T33")

    with zipfile.ZipFile(DOCX_ORIGINAL) as old_zip, zipfile.ZipFile(DOCX_AUDITADO) as new_zip:
        check(old_zip.testzip() is None and new_zip.testzip() is None, "Algún DOCX tiene CRC inválido")
        check(set(old_zip.namelist()) == set(new_zip.namelist()), "Cambió el conjunto de componentes del DOCX")
        media = [name for name in old_zip.namelist() if "/media/" in name]
        check(len(media) == 33, "Número inesperado de medios en DOCX")
        for name in media:
            check(old_zip.read(name) == new_zip.read(name), f"Cambió el medio DOCX {name}")


def sheet_paths(members: dict[str, bytes]) -> dict[str, str]:
    workbook = ET.fromstring(members["xl/workbook.xml"])
    relationships = ET.fromstring(members["xl/_rels/workbook.xml.rels"])
    targets = {
        rel.attrib["Id"]: rel.attrib["Target"]
        for rel in relationships.findall(f"{{{PKG_REL_NS}}}Relationship")
    }
    result = {}
    for sheet in workbook.find(f"{{{S_NS}}}sheets"):
        target = targets[sheet.attrib[f"{{{R_NS}}}id"]].lstrip("/")
        if not target.startswith("xl/"):
            target = "xl/" + target
        result[sheet.attrib["name"]] = target
    return result


def cells(xml: bytes) -> dict[str, tuple[dict[str, str], str | None, str | None, str | None]]:
    root = ET.fromstring(xml)
    result = {}
    for cell in root.findall(f".//{{{S_NS}}}c"):
        formula = cell.find(f"{{{S_NS}}}f")
        value = cell.find(f"{{{S_NS}}}v")
        inline = cell.find(f"{{{S_NS}}}is")
        result[cell.attrib["r"]] = (
            dict(cell.attrib),
            formula.text if formula is not None else None,
            value.text if value is not None else None,
            "".join(inline.itertext()) if inline is not None else None,
        )
    return result


def verify_xlsm() -> None:
    with zipfile.ZipFile(XLSM_ORIGINAL) as old_zip, zipfile.ZipFile(XLSM_AUDITADO) as new_zip:
        check(old_zip.testzip() is None and new_zip.testzip() is None, "Algún XLSM tiene CRC inválido")
        check(set(old_zip.namelist()) == set(new_zip.namelist()), "Cambió el conjunto de componentes del XLSM")
        old_members = {name: old_zip.read(name) for name in old_zip.namelist()}
        new_members = {name: new_zip.read(name) for name in new_zip.namelist()}

    paths = sheet_paths(old_members)
    changed_allowed = {"xl/workbook.xml", paths["Estudio Financiero Det"], paths["Modelo Riesgo Operativo"]}
    actual_changed = {name for name in old_members if old_members[name] != new_members[name]}
    check(actual_changed == changed_allowed, f"Componentes XLSM modificados inesperadamente: {actual_changed ^ changed_allowed}")

    media = [name for name in old_members if "/media/" in name]
    check(len(media) == 25, "Número inesperado de medios en XLSM")
    for name in media:
        check(old_members[name] == new_members[name], f"Cambió el medio XLSM {name}")
    binaries = [name for name in old_members if name.endswith(".bin")]
    check(not binaries, "El original contiene binarios no previstos")

    old_financial = cells(old_members[paths["Estudio Financiero Det"]])
    new_financial = cells(new_members[paths["Estudio Financiero Det"]])
    check(set(old_financial) == set(new_financial), "Cambió el conjunto de celdas de Estudio Financiero Det")
    changed_financial = {ref for ref in old_financial if old_financial[ref] != new_financial[ref]}
    check(changed_financial == {"H93", "B153", "B166"}, f"Celdas financieras modificadas inesperadamente: {changed_financial}")
    check(new_financial["H93"][1] == 'IF(ROUND(H90,2)>0,D35*$B$9/(1-(1+$B$9)^-5),0)', "Fórmula H93 incorrecta")
    check("COUNTIF(C122:K122" in (new_financial["B153"][1] or ""), "B153 aún cuenta Año 0")
    check("COUNTIF(C143:K143" in (new_financial["B166"][1] or ""), "B166 aún cuenta Año 0")
    check(math.isclose(float(new_financial["B153"][2]), 3.6712363740691614), "Caché B153 incorrecta")
    check(math.isclose(float(new_financial["B166"][2]), 3.6592994159444396), "Caché B166 incorrecta")

    old_risk = cells(old_members[paths["Modelo Riesgo Operativo"]])
    new_risk = cells(new_members[paths["Modelo Riesgo Operativo"]])
    check(set(old_risk) == set(new_risk), "Cambió el conjunto de celdas del modelo de riesgos")
    expected_risk_changes = {f"{col}{row}" for row in range(17, 28) for col in "NOPQRSTUVW"}
    changed_risk = {ref for ref in old_risk if old_risk[ref] != new_risk[ref]}
    check(changed_risk == expected_risk_changes, "Se modificaron celdas de riesgo fuera de N17:W27")
    for row in range(17, 28):
        probability_row = row - 14
        for col in "NOPQRSTUVW":
            expected = f"+_xll.RiskBinomial(1,{col}{probability_row})"
            check(new_risk[f"{col}{row}"][1] == expected, f"Fórmula de riesgo incorrecta en {col}{row}")

    workbook = ET.fromstring(new_members["xl/workbook.xml"])
    calc_pr = workbook.find(f"{{{S_NS}}}calcPr")
    check(calc_pr is not None, "Falta calcPr")
    check(calc_pr.attrib.get("calcMode") == "auto", "El cálculo no está en modo automático")
    check(calc_pr.attrib.get("fullCalcOnLoad") == "1", "No se solicitó recálculo completo")
    check(calc_pr.attrib.get("forceFullCalc") == "1", "No se forzó el recálculo")

    workbook_read = openpyxl.load_workbook(XLSM_AUDITADO, read_only=True, data_only=False, keep_vba=True)
    check("Estudio Financiero Det" in workbook_read.sheetnames, "openpyxl no abre la hoja financiera")
    check("Modelo Riesgo Operativo" in workbook_read.sheetnames, "openpyxl no abre la hoja de riesgos")
    check(workbook_read["Estudio Financiero Det"]["H93"].value.startswith("=IF(ROUND(H90,2)"), "openpyxl no interpreta H93")
    workbook_read.close()


def main() -> None:
    for path in (DOCX_ORIGINAL, DOCX_AUDITADO, XLSM_ORIGINAL, XLSM_AUDITADO):
        check(path.exists() and path.stat().st_size > 0, f"Falta {path.name}")
    verify_docx()
    verify_xlsm()
    print(f"VERIFICACIÓN SUPERADA: {CHECKS} comprobaciones")
    print(f"SHA256 DOCX original:  {digest(DOCX_ORIGINAL)}")
    print(f"SHA256 DOCX auditado:  {digest(DOCX_AUDITADO)}")
    print(f"SHA256 XLSM original:  {digest(XLSM_ORIGINAL)}")
    print(f"SHA256 XLSM auditado:  {digest(XLSM_AUDITADO)}")


if __name__ == "__main__":
    main()
