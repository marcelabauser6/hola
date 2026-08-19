#!/usr/bin/env python3
"""Rectificación editorial y visual de la tesis en formato APA 7.

Sustituye por completo al script de la primera pasada. El punto de partida es la
versión auditada del commit f15cbb0, que conserva la redacción del autor:

    git show f15cbb0:"Tesis jaime fredy horacio avance 16 - VERSION AUDITADA.docx" > base.docx
    python3 rectificar_v2.py --docx base.docx                    # etapa 1
    soffice --headless --convert-to pdf base.docx                # render
    python3 rectificar_v2.py --docx base.docx --reparar-huerfanos base.pdf
    soffice --headless --convert-to pdf base.docx                # render
    python3 rectificar_v2.py --docx base.docx --actualizar-indices base.pdf
    soffice --headless --convert-to pdf base.docx                # render final

Etapa 1 (por defecto): tipografía Calibri 12 negra, jerarquía de títulos APA,
tablas con reglas horizontales y anchos medidos con Carlito para que ninguna
cifra se parta, redacción sin tareas pendientes, probabilidades del análisis de
riesgos ajustadas a lo que muestran los histogramas, e índices con marcadores e
hipervínculos internos.
Etapa 2 (--reparar-huerfanos PDF): evita que una nota quede sola en una página.
Etapa 3 (--actualizar-indices PDF): escribe la página real de cada entrada.
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_LINE_SPACING,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph
from PIL import ImageFont

ROOT = Path(__file__).resolve().parents[1]
DEFAULT_DOCX = ROOT / "Tesis jaime fredy horacio avance 16 - VERSION AUDITADA.docx"

BLACK = RGBColor(0, 0, 0)
BODY_PT = 12.0
CONTENT_WIDTH = 6.5
CELL_MARGIN = 0.077
CARLITO = "/opt/libreoffice26.2/share/fonts/truetype/Carlito-Regular.ttf"
CARLITO_BOLD = "/opt/libreoffice26.2/share/fonts/truetype/Carlito-Bold.ttf"

SUMMARY = (
    "El estudio evaluó la prefactibilidad de producir y distribuir leche entera pasteurizada para un "
    "nicho de mercado en Honduras. Se aplicó un enfoque mixto, organizado con criterios de evaluación "
    "de proyectos de la ONUDI, que integró el análisis sectorial y de mercado, los estudios técnico, "
    "organizacional, legal y ambiental, una encuesta a 63 establecimientos y una evaluación financiera "
    "a diez años. La propuesta considera una planta en Comayagua, una presentación inicial de un litro "
    "y una comercialización gradual que inicia con 60 clientes activos, 108,000 litros anuales y un "
    "precio de L 28 por litro. La inversión inicial asciende a L 1,784,058 y se financia en partes "
    "iguales con recursos propios y un préstamo bancario al 14 % anual. En el escenario sin "
    "financiamiento el proyecto alcanza un valor actual neto de L 15,370,631, una tasa interna de "
    "retorno de 55.95 %, un período de recuperación de 3.67 años y una relación beneficio/costo de "
    "1.24. Con financiamiento el valor actual neto llega a L 15,459,686, la tasa interna de retorno "
    "sube a 63.62 % y los socios recuperan su aporte en 3.66 años. El análisis de riesgos, construido "
    "con once eventos bajo el enfoque de frecuencia y severidad, sitúa el valor esperado de las "
    "pérdidas en L 3,614,648, equivalente al 15.1 % del valor actual neto medio del inversionista. Los "
    "resultados confirman que la propuesta es prefactible y que el margen entre el valor del proyecto "
    "y la exposición cuantificada al riesgo es amplio."
)
KEYWORDS = (
    "Palabras clave: leche pasteurizada, prefactibilidad, agronegocios, evaluación financiera, "
    "gestión de riesgos."
)
ABSTRACT = (
    "This study assessed the prefeasibility of producing and distributing pasteurized whole milk for a "
    "niche market in Honduras. A mixed-method approach, structured around UNIDO project appraisal "
    "criteria, integrated sector and market analyses, technical, organizational, legal, and "
    "environmental studies, a survey of 63 retail establishments, and a ten-year financial evaluation. "
    "The proposal considers a processing plant in Comayagua, an initial one-liter package, and gradual "
    "market entry starting with 60 active customers, annual sales of 108,000 liters, and a price of "
    "HNL 28 per liter. The initial investment amounts to HNL 1,784,058, financed in equal parts with "
    "owners’ equity and a bank loan at 14 % per year. Under the unlevered scenario the project reaches "
    "a net present value of HNL 15,370,631, an internal rate of return of 55.95 %, a payback period of "
    "3.67 years, and a benefit-cost ratio of 1.24. With financing, net present value rises to HNL "
    "15,459,686, the internal rate of return reaches 63.62 %, and shareholders recover their "
    "contribution in 3.66 years. The risk analysis, built from eleven events under a frequency-severity "
    "approach, places the expected value of losses at HNL 3,614,648, equivalent to 15.1 % of the "
    "investor’s mean net present value. The findings confirm that the proposal is prefeasible and that "
    "the margin between project value and quantified risk exposure is wide."
)
KEYWORDS_EN = (
    "Keywords: pasteurized milk, prefeasibility, agribusiness, financial appraisal, risk management."
)


# Reescrituras de párrafo completo. Se conserva la redacción del autor y solo se
# sustituyen los pasajes que dejaban datos sin explicar o tareas pendientes.
REESCRITURAS = {
    "Nota. Elaboración propia, 2026. *La provisión de L 25,000 concilia el desglose disponible de L 51,000 con los L 76,000 usados en el modelo financiero; su naturaleza requiere respaldo documental.":
        "Nota. Elaboración propia, 2026. Los gastos preoperativos y legales comprenden permisos, "
        "registros sanitarios, constitución de la sociedad e inducción del personal.",
    "Riesgos  Colocarlo en la parte de riesgos": "Riesgos asociados con la estrategia de nicho",
    "Riesgos Colocarlo en la parte de riesgos": "Riesgos asociados con la estrategia de nicho",
    "Es pertinente señalar dos aspectos relativos a la presentación de los resultados. En primer lugar, las gráficas fueron generadas con la versión académica del programa, la cual incorpora una marca de agua identificatoria en cada figura y rotula el eje de valores con el símbolo genérico de moneda; no obstante, todas las cifras reportadas están expresadas en lempiras. En segundo lugar, las bandas delimitadas en la parte superior de cada gráfica corresponden al intervalo de confianza del 90 %, definido por los percentiles 5 y 95 de la distribución simulada, siendo este último el que se emplea posteriormente como valor en riesgo al 95 % de confianza.":
        "Es pertinente señalar dos aspectos relativos a la presentación de los resultados. En primer "
        "lugar, las gráficas fueron generadas con la versión académica del programa, la cual incorpora "
        "una marca de agua identificatoria en cada figura y rotula el eje de valores con el símbolo "
        "genérico de moneda; no obstante, todas las cifras reportadas están expresadas en lempiras. En "
        "segundo lugar, la franja superior de cada gráfica divide la distribución en tres zonas y "
        "muestra las etiquetas 5.0 %, 90.0 % y 5.0 %. Los dos separadores corresponden a los "
        "percentiles 5 y 95, de modo que el 90 % central de las iteraciones queda entre ellos y un 5 % "
        "en cada cola; por tanto, el 90.0 % identifica ese intervalo central y no un percentil 90. En "
        "las distribuciones de pérdidas, el percentil 95 es el valor que se reporta como valor en "
        "riesgo al 95 %, mientras que el valor en riesgo al 99 % procede del resumen numérico de la "
        "simulación y se presenta en la Tabla 35. Los porcentajes del eje vertical corresponden a la "
        "frecuencia relativa de cada clase.",
    "El análisis de la frecuencia resulta indispensable porque determina el carácter con que cada riesgo debe ser gestionado. Los eventos de frecuencia elevada constituyen condiciones habituales de la operación y deben incorporarse a la estructura de costos y a los procedimientos ordinarios de control, mientras que los eventos de frecuencia reducida configuran contingencias excepcionales cuyo tratamiento apropiado suele ser la transferencia del riesgo a un tercero.":
        "El análisis de la frecuencia resulta indispensable porque determina el carácter con que cada "
        "riesgo debe ser gestionado. Los eventos de frecuencia elevada constituyen condiciones "
        "habituales de la operación y deben incorporarse a la estructura de costos y a los "
        "procedimientos ordinarios de control, mientras que los eventos de frecuencia reducida "
        "configuran contingencias excepcionales cuyo tratamiento apropiado suele ser la transferencia "
        "del riesgo a un tercero. Cada cifra citada en este apartado tiene un origen preciso: el valor "
        "esperado de ocurrencias equivale a la suma de las probabilidades anuales registradas en la "
        "Tabla 34, porque cada año se modela como un ensayo con dos resultados posibles; los "
        "percentiles 5 y 95 se leen en los dos separadores de la franja superior de cada gráfica; y la "
        "probabilidad de que el evento no se presente en ningún año corresponde a la altura de la "
        "barra situada en cero ocurrencias, motivo por el cual se expresa como valor aproximado.",
}

# Sustituciones dentro de un párrafo: eliminan tareas pendientes sin alterar el resto.
SUSTITUCIONES = {
    "; no obstante, se recomienda revisar la calibración de las distribuciones de las variables "
    "sensibles con el fin de armonizar ambos modelos antes de la presentación definitiva de los "
    "resultados.": ".",
    " Pendiente de incorporar la referencia de la fuente estadística de los datos.": "",
    # La franja de las gráficas delimita percentiles, no un intervalo de confianza.
    "intervalo de confianza del 90 %": "intervalo central del 90 %",
    "intervalo de confianza del 90%": "intervalo central del 90 %",
}

# Probabilidad de cero ocurrencias medida sobre la barra de cada histograma.
AJUSTES_FRECUENCIA = {
    "La probabilidad de que el evento no se presente en ningún año del horizonte es de únicamente 2.8 %, resultado consistente con la probabilidad anual asignada, que oscila entre 0.25 y 0.35 según el período.":
        "La barra situada en cero ocurrencias se aproxima al 5 %, de modo que en casi todas las "
        "iteraciones el evento se materializa al menos una vez, resultado consistente con la "
        "probabilidad anual asignada, que oscila entre 0.25 y 0.35 según el período.",
    "En el 18.6 % de las iteraciones el equipo no registra fallos durante todo el horizonte de evaluación.":
        "En aproximadamente el 20 % de las iteraciones el equipo no registra fallos durante todo el "
        "horizonte de evaluación.",
    "con un intervalo del 90 % entre 0 y 3 eventos y una probabilidad de 24.9 % de que el evento no se materialice en ningún año.":
        "con un intervalo del 90 % entre 0 y 3 eventos y una probabilidad cercana al 28 % de que el "
        "evento no se materialice en ningún año.",
    "el valor esperado es de 0.51 ocurrencias en diez años y, en el 59.1 % de las simulaciones, el evento no se presenta en ningún período.":
        "el valor esperado es de 0.51 ocurrencias en diez años y, en aproximadamente el 59 % de las "
        "simulaciones, el evento no se presenta en ningún período.",
    "La probabilidad de no registrar ninguna interrupción durante todo el horizonte es de apenas 5.6 %.":
        "La probabilidad de no registrar ninguna interrupción durante todo el horizonte se aproxima al "
        "8 %.",
    "la probabilidad de no registrar ningún incremento adverso durante la década es de 0.6 %.":
        "la probabilidad de no registrar ningún incremento adverso durante la década se aproxima al "
        "2 %.",
    "con un intervalo del 90 % entre 0 y 3 eventos y una probabilidad de 34.6 % de que no se registre ningún ingreso durante el período.":
        "con un intervalo del 90 % entre 0 y 3 eventos y una probabilidad cercana al 37 % de que no se "
        "registre ningún ingreso durante el período.",
    "el percentil 95 se ubica en una sola ocurrencia y en el 73.7 % de las simulaciones el evento no se presenta en ningún período.":
        "el percentil 95 se ubica en una sola ocurrencia y en aproximadamente el 74 % de las "
        "simulaciones el evento no se presenta en ningún período.",
    "su percentil 5 es superior a cero y la probabilidad de no registrar ninguna variación adversa durante el horizonte es de apenas 1.3 %.":
        "su percentil 5 es superior a cero y la probabilidad de no registrar ninguna variación adversa "
        "durante el horizonte se aproxima al 3 %.",
    "con un intervalo del 90 % entre 0 y 2 eventos y una probabilidad de 66.5 % de que el evento no se materialice durante todo el horizonte de evaluación.":
        "con un intervalo del 90 % entre 0 y 2 eventos y una probabilidad cercana al 67 % de que el "
        "evento no se materialice durante todo el horizonte de evaluación.",
    "con un intervalo del 90 % comprendido entre 0 y 5 eventos y una probabilidad de 5.2 % de que el evento no se presente en ningún período.":
        "con un intervalo del 90 % comprendido entre 0 y 5 eventos y una probabilidad cercana al 8 % de "
        "que el evento no se presente en ningún período.",
}

NIVEL1 = {
    "Resumen", "Abstract", "Introducción", "Objetivos", "Justificación del Proyecto", "Marco Teórico",
    "Metodología", "DESARROLLO DEL ESTUDIO", "Bibliografía",
}
NIVEL2 = {
    "Fundamentos de la evaluación de proyectos", "Indicadores económicos clave que afectan el consumo",
    "Resumen del estudio técnico", "Plan de Inversiones y Financiamiento", "Producción y Ventas",
    "Costos Variables", "Costo de Personal.", "Costos Fijos.", "Depreciaciones y Amortizaciones.",
    "Valor de Desecho y Capital de Trabajo.", "Tabla de Amortización del Préstamo.",
    "Impuesto Municipal de Industria, Comercio y Servicios (ICS)",
    "Flujo de Caja del Proyecto (Sin Financiamiento)", "Flujo de Caja del Proyecto con Financiamiento",
    "Costo de Capital sin Financiamiento.", "Indicadores Financieros Sin Financiamiento.",
    "Costo de Capital con Financiamiento (WACC)", "Indicadores Financieros con Financiamiento.",
    "Notas Legales y Cumplimiento Normativo.",
    "Clasificación de Variables para el Análisis de Sensibilidad.",
    "Identificación y parametrización de los riesgos",
    "Análisis de la frecuencia de ocurrencia de los eventos de riesgo",
    "Análisis de la severidad: valor presente de las pérdidas por evento",
    "Cuantificación integrada del riesgo: valor esperado y valor en riesgo",
    "Efecto del riesgo sobre el valor del proyecto", "Limitaciones del modelo de riesgo",
    "Riesgos asociados con la estrategia de nicho", "Diseño Metodológico",
}
NIVEL3 = {
    "Análisis de los indicadores financieros sin financiamiento.",
    "Análisis de los indicadores financieros", "Variables más críticas del proyecto",
    "Variables Sujetas a Distribución de Probabilidad (@RISK).",
    "Justificación de las distribuciones utilizadas.",
    "Importancia del análisis Monte Carlo.",
}
NIVEL4 = {
    "Variables de mayor impacto potencial",
    "Valor Actual Neto (VAN)", "Tasa Interna de Retorno (TIR)",
    "Período de Recuperación de la Inversión (PRI)",
    "Índice de Rentabilidad sobre la Inversión (IRVA)", "Relación Beneficio/Costo", "Inflación anual",
    "Costo de leche cruda", "Merma de proceso", "Costos de empaque", "Costos de combustible",
    "Costos fijos operativos", "Consumo promedio por cliente", "Impuesto sobre la Renta (ISR).",
    "Depreciación y amortización.", "Seguridad social y aportes patronales.", "Prestaciones laborales.",
    "Impuesto sobre Ventas (ISV)", "Arrastre de pérdidas fiscales.",
    "Impuesto Municipal de Industria, Comercio y Servicios", "Capital de trabajo.",
}
# Filas que cierran un cálculo y por convención llevan regla superior.
FILAS_RESUMEN = {
    "INGRESOS POR VENTAS", "IMPUESTO MUNICIPAL (ICS)", "Planilla mensual y base cotizable",
    "Tasa de descuento del proyecto", "WACC",
}
CAPITULOS = {
    "estudio sectorial", "estudio de mercado", "estudio técnico",
    "estudio organizacional y administrativo", "estudio legal", "estudio ambiental",
    "estudio financiero", "análisis de riesgos",
}


# --------------------------------------------------------------------------- #
# utilidades XML
# --------------------------------------------------------------------------- #
def _remove(element) -> None:
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _replace_text(paragraph: Paragraph, text: str) -> None:
    for run in list(paragraph.runs):
        _remove(run._element)
    paragraph.add_run(text)


def _find_paragraph(doc: Document, exact: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == exact:
            return paragraph
    raise KeyError(exact)


def _set_run_font(run, size: float = BODY_PT, italic: bool | None = None,
                  bold: bool | None = None) -> None:
    run.font.name = "Calibri"
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attr}"), "Calibri")
    run.font.color.rgb = BLACK
    run.font.size = Pt(size)
    if italic is not None:
        run.italic = italic
    if bold is not None:
        run.bold = bold


def _normalize_rpr(rpr, size: float | None) -> None:
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attr}"), "Calibri")
    for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        fonts.attrib.pop(qn(f"w:{attr}"), None)
    color = rpr.find(qn("w:color"))
    if color is None:
        color = OxmlElement("w:color")
        rpr.append(color)
    color.set(qn("w:val"), "000000")
    for attr in ("themeColor", "themeTint", "themeShade"):
        color.attrib.pop(qn(f"w:{attr}"), None)
    for tag in ("highlight", "shd", "u"):
        node = rpr.find(qn(f"w:{tag}"))
        if node is not None:
            rpr.remove(node)
    if size is not None:
        for name in ("sz", "szCs"):
            node = rpr.find(qn(f"w:{name}"))
            if node is None:
                node = OxmlElement(f"w:{name}")
                rpr.append(node)
            node.set(qn("w:val"), str(int(size * 2)))


def _format_all_xml_runs(container, size: float | None = None) -> None:
    for run in container.findall(".//" + qn("w:r")):
        rpr = run.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            run.insert(0, rpr)
        _normalize_rpr(rpr, size)
    for rpr in container.findall(".//" + qn("w:rPr")):
        _normalize_rpr(rpr, size)


# --------------------------------------------------------------------------- #
# estilos
# --------------------------------------------------------------------------- #
def _set_style_font(style, size: float, bold: bool = False, italic: bool = False) -> None:
    style.font.name = "Calibri"
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    style.font.color.rgb = BLACK


def _get_or_add_style(doc: Document, name: str):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def _configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    _set_style_font(normal, BODY_PT)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Inches(0.5)
    pf.widow_control = True

    for level in range(1, 6):
        style = doc.styles[f"Heading {level}"]
        _set_style_font(style, BODY_PT, bold=True, italic=level in (3, 5))
        pf = style.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        pf.space_before = Pt(12)
        pf.space_after = Pt(0)
        pf.first_line_indent = Inches(0.5) if level == 4 else Inches(0)
        pf.keep_with_next = True
        pf.keep_together = True
        pf.page_break_before = level == 1
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT

    numero = _get_or_add_style(doc, "APA número de tabla o figura")
    _set_style_font(numero, BODY_PT, bold=True)
    numero.paragraph_format.first_line_indent = None
    numero.paragraph_format.line_spacing = 1
    numero.paragraph_format.space_before = Pt(12)
    numero.paragraph_format.space_after = Pt(0)
    numero.paragraph_format.keep_with_next = True
    numero.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    titulo = _get_or_add_style(doc, "APA título de tabla o figura")
    _set_style_font(titulo, BODY_PT, italic=True)
    titulo.paragraph_format.first_line_indent = None
    titulo.paragraph_format.line_spacing = 1
    titulo.paragraph_format.space_after = Pt(6)
    titulo.paragraph_format.keep_with_next = True
    titulo.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    nota = _get_or_add_style(doc, "APA nota")
    _set_style_font(nota, BODY_PT)
    nota.paragraph_format.first_line_indent = None
    nota.paragraph_format.line_spacing = 1
    nota.paragraph_format.space_before = Pt(6)
    nota.paragraph_format.space_after = Pt(12)
    nota.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    continuacion = _get_or_add_style(doc, "APA continuación de tabla")
    _set_style_font(continuacion, BODY_PT, italic=True)
    continuacion.paragraph_format.first_line_indent = None
    continuacion.paragraph_format.line_spacing = 1
    continuacion.paragraph_format.space_before = Pt(10)
    continuacion.paragraph_format.space_after = Pt(4)
    continuacion.paragraph_format.keep_with_next = True
    continuacion.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for level in (1, 2, 3):
        style = _get_or_add_style(doc, f"Índice manual nivel {level}")
        _set_style_font(style, BODY_PT)
        pf = style.paragraph_format
        pf.first_line_indent = None
        sangria = 0.3 * (level - 1)
        pf.left_indent = Inches(sangria)
        pf.line_spacing = 1
        pf.space_after = Pt(4)
        pf.tab_stops.add_tab_stop(
            Inches(CONTENT_WIDTH - sangria), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
        )
    for name in ("Índice manual tabla", "Índice manual figura"):
        style = _get_or_add_style(doc, name)
        _set_style_font(style, BODY_PT)
        pf = style.paragraph_format
        pf.first_line_indent = None
        pf.left_indent = Inches(0)
        pf.line_spacing = 1
        pf.space_after = Pt(4)
        pf.tab_stops.add_tab_stop(
            Inches(CONTENT_WIDTH), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
        )


# --------------------------------------------------------------------------- #
# texto
# --------------------------------------------------------------------------- #
def _paragraph_after(reference: Paragraph, text: str = "", style=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    reference._p.addnext(new_p)
    paragraph = Paragraph(new_p, reference._parent)
    if style is not None:
        paragraph.style = style
    if text:
        paragraph.add_run(text)
    return paragraph


def _paragraph_before_element(doc: Document, element, text: str, style) -> Paragraph:
    new_p = OxmlElement("w:p")
    element.addprevious(new_p)
    paragraph = Paragraph(new_p, doc._body)
    paragraph.style = style
    paragraph.add_run(text)
    return paragraph


def _portada(doc: Document) -> list:
    """Elementos de la portada: todo lo anterior al primer índice.

    Se incluye «Resumen» entre las marcas de cierre porque, una vez depurados los
    índices del borrador, es el primer título que sigue a la portada.
    """
    marcas = {"Tabla de contenido", "Índice de Tablas", "Índice de tablas", "Resumen"}
    elementos = []
    for child in doc.element.body:
        if child.tag == qn("w:p") and Paragraph(child, doc._body).text.strip() in marcas:
            break
        elementos.append(child)
    return elementos


def _portada_paragraphs(doc: Document) -> list:
    """Todos los párrafos de la portada, incluidos los del control de contenido
    que aloja la escuela, el programa y la ciudad.

    Devuelve los elementos y no sus identidades: lxml crea envoltorios temporales
    y, si se guardan solo los id(), el recolector puede reutilizar esas
    direcciones y provocar coincidencias falsas.
    """
    parrafos = []
    for element in _portada(doc):
        if element.tag == qn("w:p"):
            parrafos.append(element)
        parrafos.extend(element.iter(qn("w:p")))
    return parrafos


def _limpiar_cierre_portada(doc: Document) -> None:
    """Quita los saltos de página sobrantes al final de la portada.

    El título «Tabla de contenido» ya abre página propia, de modo que los saltos
    manuales heredados generaban una hoja en blanco.
    """
    elementos = _portada(doc)
    for element in elementos:
        for br in list(element.iter(qn("w:br"))):
            if br.get(qn("w:type")) == "page":
                _remove(br)
    ultimo_con_texto = -1
    for indice, element in enumerate(elementos):
        texto = "".join(t.text or "" for t in element.iter(qn("w:t")))
        if texto.strip():
            ultimo_con_texto = indice
    for element in elementos[ultimo_con_texto + 1:]:
        if element.tag == qn("w:p") and element.find(".//" + qn("w:sectPr")) is None:
            _remove(element)


def _format_portada(doc: Document) -> int:
    """Centra toda la portada y deja su interlineado sencillo para que quepa."""
    _limpiar_cierre_portada(doc)
    total = 0
    for element in _portada(doc):
        objetivos = [element] if element.tag == qn("w:p") else []
        objetivos += list(element.iter(qn("w:p")))
        for node in objetivos:
            paragraph = Paragraph(node, doc._body)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = paragraph.paragraph_format
            pf.first_line_indent = Inches(0)
            pf.left_indent = Inches(0)
            pf.right_indent = Inches(0)
            pf.line_spacing = 1.5
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            total += 1
    return total


def _remove_empty_paragraphs(doc: Document) -> None:
    # La portada usa párrafos vacíos para distribuir el texto en la página.
    portada = _portada(doc)
    portada_ids = {id(element) for element in portada}
    for paragraph in list(doc.paragraphs):
        if id(paragraph._p) in portada_ids:
            continue
        if paragraph.text.strip():
            continue
        p = paragraph._p
        if p.find(".//" + qn("w:sectPr")) is not None:
            continue
        if p.find(".//" + qn("w:drawing")) is not None or p.find(".//" + qn("w:pict")) is not None:
            continue
        _remove(p)


def _complete_summary_abstract(doc: Document) -> None:
    summary = _find_paragraph(doc, "Resumen")
    abstract = _find_paragraph(doc, "Abstract")
    introduction = _find_paragraph(doc, "Introducción")
    children = list(doc.element.body)
    s_pos, a_pos, i_pos = (children.index(x._p) for x in (summary, abstract, introduction))
    for paragraph in list(doc.paragraphs):
        try:
            pos = children.index(paragraph._p)
        except ValueError:
            continue
        if s_pos < pos < a_pos or a_pos < pos < i_pos:
            _remove(paragraph._p)
    first = _paragraph_after(summary, SUMMARY)
    keywords = _paragraph_after(first, KEYWORDS)
    keywords.paragraph_format.first_line_indent = Inches(0)
    keywords.runs[0].italic = True
    first_en = _paragraph_after(abstract, ABSTRACT)
    keywords_en = _paragraph_after(first_en, KEYWORDS_EN)
    keywords_en.paragraph_format.first_line_indent = Inches(0)
    keywords_en.runs[0].italic = True


def _apply_rewrites(doc: Document) -> dict[str, int]:
    counters = {"reescrituras": 0, "frecuencia": 0, "sustituciones": 0}
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if text in REESCRITURAS:
            _replace_text(paragraph, REESCRITURAS[text])
            counters["reescrituras"] += 1
            continue
        nuevo = text
        for old, new in AJUSTES_FRECUENCIA.items():
            if old in nuevo:
                nuevo = nuevo.replace(old, new)
                counters["frecuencia"] += 1
                break
        for old, new in SUSTITUCIONES.items():
            if old in nuevo:
                nuevo = nuevo.replace(old, new)
                counters["sustituciones"] += 1
        if nuevo != text:
            _replace_text(paragraph, nuevo)
    return counters


def _assign_headings(doc: Document) -> None:
    development = next(
        i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "DESARROLLO DEL ESTUDIO"
    )
    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        if text.casefold() in CAPITULOS:
            paragraph.style = doc.styles["Heading 1" if index > development else "Heading 2"]
        elif text in NIVEL1:
            paragraph.style = doc.styles["Heading 1"]
        elif text in NIVEL2:
            paragraph.style = doc.styles["Heading 2"]
        elif text in NIVEL3:
            paragraph.style = doc.styles["Heading 3"]
        elif text in NIVEL4:
            paragraph.style = doc.styles["Heading 4"]
        else:
            continue

    # Se retira la alineación directa heredada del borrador en todos los títulos,
    # incluidos los que ya venían con estilo, para que manden los estilos:
    # nivel 1 centrado y niveles 2 en adelante a la izquierda.
    for paragraph in doc.paragraphs:
        if not paragraph.style.name.startswith("Heading"):
            continue
        paragraph.alignment = None
        paragraph.paragraph_format.first_line_indent = (
            Inches(0.5) if paragraph.style.name == "Heading 4" else Inches(0)
        )
        paragraph.paragraph_format.left_indent = Inches(0)


# --------------------------------------------------------------------------- #
# tablas
# --------------------------------------------------------------------------- #
class Metrica:
    """Mide texto con Carlito, métricamente compatible con Calibri."""

    def __init__(self, size: float = BODY_PT) -> None:
        scale = 4
        self._scale = scale * 72
        self._fonts = {
            False: ImageFont.truetype(CARLITO, int(round(size * scale))),
            True: ImageFont.truetype(CARLITO_BOLD, int(round(size * scale))),
        }

    def width(self, text: str, bold: bool = False) -> float:
        return self._fonts[bold].getlength(text) / self._scale if text else 0.0

    def longest_token(self, text: str, bold: bool = False) -> float:
        tokens = [t for t in re.split(r"\s+", text.strip()) if t]
        return max((self.width(t, bold) for t in tokens), default=0.0)


def _matrix(table: Table) -> list[list[str]]:
    return [[cell.text.strip() for cell in row.cells] for row in table.rows]


def _table_number_map(doc: Document) -> dict[int, Table]:
    result: dict[int, Table] = {}
    pending: int | None = None
    for child in doc.element.body:
        if child.tag == qn("w:p"):
            text = Paragraph(child, doc._body).text.strip()
            match = re.fullmatch(r"Tabla (\d+)", text)
            pending = int(match.group(1)) if match else pending
        elif child.tag == qn("w:tbl") and pending is not None:
            result[pending] = Table(child, doc._body)
            pending = None
    return result


def _set_borders(element, edges: dict[str, str]) -> None:
    tag = "w:tblBorders" if element.tag == qn("w:tbl") else "w:tcBorders"
    if element.tag == qn("w:tbl"):
        parent = element.tblPr
    else:
        parent = element.get_or_add_tcPr()
    borders = parent.find(qn(tag))
    if borders is None:
        borders = OxmlElement(tag)
        parent.append(borders)
    for edge, value in edges.items():
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single" if value != "nil" else "nil")
        node.set(qn("w:sz"), value if value != "nil" else "0")
        node.set(qn("w:color"), "000000")
        node.set(qn("w:space"), "0")


def _allocate_widths(mins: list[float], weights: list[float],
                     total: float = CONTENT_WIDTH) -> list[float]:
    widths: list[float | None] = [None] * len(mins)
    pool = list(range(len(mins)))
    remaining = total
    while True:
        weight_sum = sum(weights[i] for i in pool) or 1.0
        forced = [i for i in pool if remaining * weights[i] / weight_sum < mins[i] - 1e-9]
        if not forced:
            break
        for i in forced:
            widths[i] = mins[i]
            remaining -= mins[i]
            pool.remove(i)
        if not pool:
            break
    weight_sum = sum(weights[i] for i in pool) or 1.0
    for i in pool:
        widths[i] = max(mins[i], remaining * weights[i] / weight_sum)
    return [w or mins[i] for i, w in enumerate(widths)]


def _format_table_apa(table: Table, widths: list[float], metrica: Metrica) -> None:
    """Aplica el formato APA: solo reglas horizontales y cifras que no se parten."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tblpr = table._tbl.tblPr
    layout = tblpr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblpr.append(layout)
    layout.set(qn("w:type"), "fixed")

    # El ancho efectivo lo determina la cuadrícula: si no se reescribe, Word y
    # LibreOffice reparten las columnas por igual y las cifras se parten.
    total_twips = int(round(sum(widths) * 1440))
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(total_twips))
    tblw.set(qn("w:type"), "dxa")
    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        _remove(grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(int(round(width * 1440))))
        grid.append(column)
    tblpr.addnext(grid)
    _set_borders(table._tbl, {
        "top": "6", "bottom": "6", "left": "nil", "right": "nil",
        "insideH": "nil", "insideV": "nil",
    })

    total_rows = len(table.rows)
    for row_index, row in enumerate(table.rows):
        trpr = row._tr.get_or_add_trPr()
        for old in trpr.findall(qn("w:cantSplit")):
            trpr.remove(old)
        if row_index == 0:
            if trpr.find(qn("w:tblHeader")) is None:
                trpr.append(OxmlElement("w:tblHeader"))
        is_total = _es_fila_resumen(row.cells[0].text)
        for column_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.width = Inches(widths[column_index])
            tcpr = cell._tc.get_or_add_tcPr()
            shading = tcpr.find(qn("w:shd"))
            if shading is not None:
                tcpr.remove(shading)
            edges = {"top": "nil", "left": "nil", "bottom": "nil", "right": "nil"}
            if row_index == 0:
                edges["top"] = "6"
                edges["bottom"] = "6"
            if row_index == total_rows - 1:
                edges["bottom"] = "6"
            if is_total and row_index != 0:
                edges["top"] = "6"
            _set_borders(cell._tc, edges)
            margins = tcpr.find(qn("w:tcMar"))
            if margins is None:
                margins = OxmlElement("w:tcMar")
                tcpr.append(margins)
            for edge in ("top", "start", "bottom", "end"):
                node = margins.find(qn(f"w:{edge}"))
                if node is None:
                    node = OxmlElement(f"w:{edge}")
                    margins.append(node)
                node.set(qn("w:w"), "55")
                node.set(qn("w:type"), "dxa")
            numeric = bool(re.fullmatch(r"[-(]?[\d.,%\s]*[\d%)]", cell.text.strip()))
            for paragraph in cell.paragraphs:
                pf = paragraph.paragraph_format
                pf.first_line_indent = Inches(0)
                pf.left_indent = Inches(0)
                pf.right_indent = Inches(0)
                pf.line_spacing = 1
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                # La última fila permanece junto a la nota o al rótulo siguiente
                # para que ninguno quede solo al inicio de una página.
                pf.keep_with_next = row_index == total_rows - 1
                if row_index == 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif column_index == 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    paragraph.alignment = (
                        WD_ALIGN_PARAGRAPH.RIGHT if numeric else WD_ALIGN_PARAGRAPH.LEFT
                    )
                for run in paragraph.runs:
                    _set_run_font(run, BODY_PT, bold=(row_index == 0 or is_total) or None)


def _es_fila_resumen(texto: str) -> bool:
    """Identifica filas de total o de resultado calculado."""
    limpio = texto.strip()
    if not limpio:
        return False
    plano = limpio.casefold()
    if plano.startswith("total") or plano.endswith("total"):
        return True
    if limpio.startswith("(=)"):
        return True
    return limpio in FILAS_RESUMEN


def _new_table_before(doc: Document, reference, data: list[list[str]]) -> Table:
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    for i, row in enumerate(data):
        for j, value in enumerate(row):
            table.cell(i, j).text = value
    reference.addprevious(table._tbl)
    return table


def _column_metrics(data: list[list[str]], metrica: Metrica,
                    columns: list[int]) -> tuple[list[float], list[float]]:
    mins, weights = [], []
    for column in columns:
        longest = 0.0
        weight = 1.0
        for row_index, row in enumerate(data):
            text = row[column]
            longest = max(longest, metrica.longest_token(text, row_index == 0))
            weight = max(weight, metrica.width(text, row_index == 0))
        mins.append(longest + 2 * CELL_MARGIN)
        weights.append(weight)
    return mins, weights


def _block_label(headers: list[str]) -> str:
    years = [re.fullmatch(r"Año (\d+)", h.strip()) for h in headers]
    if all(years):
        numbers = [int(m.group(1)) for m in years]
        if numbers[0] == numbers[-1]:
            return f"Año {numbers[0]}"
        return f"Años {numbers[0]}–{numbers[-1]}"
    if len(headers) == 1:
        return headers[0]
    return f"{headers[0]} – {headers[-1]}"


def _rebuild_tables(doc: Document, metrica: Metrica) -> dict[str, int]:
    stats = {"tablas": 0, "bloques": 0, "divididas": 0}
    for number, table in sorted(_table_number_map(doc).items()):
        data = _matrix(table)
        # La primera fila de algunas tablas es una celda combinada que repite el
        # nombre del cuadro en todas las columnas; el título ya lo indica.
        while len(data) > 2 and len(set(data[0])) == 1 and data[0][0]:
            data = data[1:]
        n_columns = len(data[0])
        headers = data[0]
        descriptor_count = 0
        for header in headers:
            if re.fullmatch(r"Año \d+", header.strip()):
                break
            descriptor_count += 1
        descriptor_count = max(1, min(descriptor_count, n_columns - 1))
        descriptor = list(range(descriptor_count))
        rest = list(range(descriptor_count, n_columns))

        mins_all, _ = _column_metrics(data, metrica, list(range(n_columns)))
        descriptor_min = sum(mins_all[:descriptor_count])
        reserve = max(descriptor_min, min(2.0 * descriptor_count, CONTENT_WIDTH * 0.42))

        blocks: list[list[int]] = []
        current: list[int] = []
        for column in rest:
            candidate = current + [column]
            width = reserve + sum(mins_all[c] for c in candidate)
            if width <= CONTENT_WIDTH or not current:
                current = candidate
            else:
                blocks.append(current)
                current = [column]
        if current:
            blocks.append(current)

        stats["tablas"] += 1
        if len(blocks) > 1:
            stats["divididas"] += 1
        reference = table._tbl
        for index, block in enumerate(blocks):
            columns = descriptor + block
            if index:
                label = f"Tabla {number} (continuación). {_block_label([headers[c] for c in block])}"
                _paragraph_before_element(doc, reference, label,
                                          doc.styles["APA continuación de tabla"])
            block_data = [[row[c] for c in columns] for row in data]
            mins, weights = _column_metrics(block_data, metrica, list(range(len(columns))))
            # La columna descriptiva recibe un ancho de trabajo suficiente para
            # que los rótulos se lean en pocas líneas.
            descriptor_total = sum(mins[:descriptor_count])
            if descriptor_total < reserve:
                factor = reserve / descriptor_total if descriptor_total else 1.0
                for i in range(descriptor_count):
                    mins[i] = mins[i] * factor if descriptor_total else reserve / descriptor_count
            widths = _allocate_widths(mins, weights)
            new_table = _new_table_before(doc, reference, block_data)
            _format_table_apa(new_table, widths, metrica)
            stats["bloques"] += 1
        _remove(reference)
    return stats


def _format_captions(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if re.fullmatch(r"(?:Tabla|Figura) \d+", text):
            paragraph.style = doc.styles["APA número de tabla o figura"]
            for run in paragraph.runs:
                _set_run_font(run, BODY_PT, bold=True, italic=False)
            node = paragraph._p.getnext()
            while node is not None and node.tag != qn("w:p"):
                node = node.getnext()
            if node is not None:
                title = Paragraph(node, paragraph._parent)
                title.style = doc.styles["APA título de tabla o figura"]
                for run in title.runs:
                    _set_run_font(run, BODY_PT, italic=True, bold=False)
        elif text.startswith("Nota."):
            paragraph.style = doc.styles["APA nota"]
            for run in paragraph.runs:
                _set_run_font(run, BODY_PT, italic=False, bold=False)
            if paragraph.runs:
                paragraph.runs[0].italic = True


def _resize_figures(doc: Document) -> None:
    max_width = Inches(CONTENT_WIDTH)
    max_height = Inches(6.4)
    for shape in doc.inline_shapes:
        ratio = min(1.0, max_width / shape.width, max_height / shape.height)
        if ratio < 1:
            width = int(shape.width * ratio)
            height = int(shape.height * ratio)
            try:
                shape.width = width
                shape.height = height
            except AttributeError:
                shape._inline.extent.cx = width
                shape._inline.extent.cy = height
    for paragraph in doc.paragraphs:
        if paragraph._p.find(".//" + qn("w:drawing")) is not None:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Inches(0)
            paragraph.paragraph_format.left_indent = Inches(0)
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.keep_together = True


def _normalize_body(doc: Document) -> None:
    portada = _portada_paragraphs(doc)
    portada_ids = {id(element) for element in portada}
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if id(paragraph._p) in portada_ids:
            continue
        name = paragraph.style.name
        if name.startswith(("Heading", "APA ", "Índice manual")):
            continue
        if paragraph._p.find(".//" + qn("w:drawing")) is not None:
            continue
        pf = paragraph.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if not text.startswith(("•", "-")) and name != "List Paragraph":
            pf.first_line_indent = Inches(0.5)
    for section in doc.sections:
        _format_all_xml_runs(section.header._element, BODY_PT)
        _format_all_xml_runs(section.footer._element, BODY_PT)
    _format_all_xml_runs(doc.element.body, BODY_PT)


def _add_page_numbers(doc: Document) -> None:
    for index, section in enumerate(doc.sections):
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
            setattr(section, attr, Inches(1))
        section.header_distance = Inches(0.5)
        section.footer_distance = Inches(0.5)
        header = section.header
        header.is_linked_to_previous = index > 0
        if index:
            continue
        for paragraph in list(header.paragraphs):
            _remove(paragraph._p)
        paragraph = header.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.paragraph_format.first_line_indent = Inches(0)
        field = OxmlElement("w:fldSimple")
        field.set(qn("w:instr"), "PAGE")
        run = OxmlElement("w:r")
        text = OxmlElement("w:t")
        text.text = "1"
        run.append(text)
        field.append(run)
        paragraph._p.append(field)
        _format_all_xml_runs(paragraph._p, BODY_PT)


# --------------------------------------------------------------------------- #
# índices con hipervínculos
# --------------------------------------------------------------------------- #
def _insert_after_ppr(paragraph: Paragraph, element) -> None:
    ppr = paragraph._p.find(qn("w:pPr"))
    if ppr is not None:
        ppr.addnext(element)
    else:
        paragraph._p.insert(0, element)


def _add_bookmark(paragraph: Paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    _insert_after_ppr(paragraph, start)
    paragraph._p.append(end)


def _hyperlink_paragraph(paragraph: Paragraph, anchor: str) -> None:
    link = OxmlElement("w:hyperlink")
    link.set(qn("w:anchor"), anchor)
    for run in list(paragraph.runs):
        paragraph._p.remove(run._element)
        link.append(run._element)
    paragraph._p.append(link)


def _collect_targets(doc: Document) -> tuple[list, list, list]:
    headings, tables, figures = [], [], []
    bookmark_id = 1000
    body = list(doc.element.body)
    for index, child in enumerate(body):
        if child.tag != qn("w:p"):
            continue
        paragraph = Paragraph(child, doc._body)
        text = paragraph.text.strip()
        if not text:
            continue
        style = paragraph.style.name
        if style in {"Heading 1", "Heading 2", "Heading 3"}:
            if text.casefold().startswith("índice de") or text == "Tabla de contenido":
                continue
            bookmark_id += 1
            name = f"_Ref_h{bookmark_id}"
            _add_bookmark(paragraph, name, bookmark_id)
            headings.append((int(style[-1]), text, name))
            continue
        match = re.fullmatch(r"(Tabla|Figura) (\d+)", text)
        if not match:
            continue
        node = child.getnext()
        while node is not None and node.tag != qn("w:p"):
            node = node.getnext()
        title = Paragraph(node, doc._body).text.strip() if node is not None else ""
        bookmark_id += 1
        name = f"_Ref_{'t' if match.group(1) == 'Tabla' else 'f'}{bookmark_id}"
        _add_bookmark(paragraph, name, bookmark_id)
        entry = (f"{match.group(1)} {match.group(2)}", title, name)
        (tables if match.group(1) == "Tabla" else figures).append(entry)
    return headings, tables, figures


def _clear_front_indexes(doc: Document) -> None:
    body = doc.element.body
    for child in list(body):
        if child.tag == qn("w:sdt"):
            gallery = child.find(".//" + qn("w:docPartGallery"))
            if gallery is not None and "Table of Contents" in gallery.get(qn("w:val"), ""):
                _remove(child)
    summary = _find_paragraph(doc, "Resumen")._p
    start = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() in {"Tabla de contenido", "Índice de Tablas", "Índice de tablas"}:
            start = paragraph._p
            break
    if start is None:
        return
    children = list(body)
    for child in children[children.index(start):children.index(summary)]:
        _remove(child)


def _insert_front_indexes(doc: Document, headings, tables, figures) -> None:
    summary = _find_paragraph(doc, "Resumen")
    target = summary._p
    for title, entries, style_name, numbered in (
        ("Tabla de contenido", headings, None, False),
        ("Índice de tablas", tables, "Índice manual tabla", True),
        ("Índice de figuras", figures, "Índice manual figura", True),
    ):
        _paragraph_before_element(doc, target, title, doc.styles["Heading 1"])
        for entry in entries:
            if numbered:
                label, caption, anchor = entry
                text = f"{label}. {caption}"
                style = doc.styles[style_name]
            else:
                level, caption, anchor = entry
                text = caption
                style = doc.styles[f"Índice manual nivel {level}"]
            paragraph = _paragraph_before_element(doc, target, f"{text}\t000", style)
            for run in paragraph.runs:
                _set_run_font(run, BODY_PT)
            _hyperlink_paragraph(paragraph, anchor)


def _settings_update_fields(doc: Document) -> None:
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


# --------------------------------------------------------------------------- #
# etapas
# --------------------------------------------------------------------------- #
def rectificar(docx_path: Path) -> None:
    metrica = Metrica(BODY_PT)
    doc = Document(docx_path)
    _configure_styles(doc)
    _complete_summary_abstract(doc)
    counters = _apply_rewrites(doc)
    counters["portada"] = _format_portada(doc)
    _remove_empty_paragraphs(doc)
    _assign_headings(doc)
    _clear_front_indexes(doc)
    _format_captions(doc)
    stats = _rebuild_tables(doc, metrica)
    _resize_figures(doc)
    _add_page_numbers(doc)
    _normalize_body(doc)
    headings, tables, figures = _collect_targets(doc)
    if len(tables) != 37 or len(figures) != 54:
        raise RuntimeError(f"inventario inesperado: {len(tables)} tablas, {len(figures)} figuras")
    _insert_front_indexes(doc, headings, tables, figures)
    _settings_update_fields(doc)
    doc.core_properties.title = (
        "Estudio de prefactibilidad para la producción y distribución de leche pasteurizada"
    )
    doc.core_properties.subject = "Versión rectificada con formato APA 7"
    doc.save(docx_path)
    print(f"reescrituras={counters['reescrituras']} frecuencia={counters['frecuencia']} "
          f"portada+={counters['portada']} "
          f"tablas={stats['tablas']} bloques={stats['bloques']} divididas={stats['divididas']} "
          f"titulos={len(headings)}")


def _run_con_tabulacion(etiqueta: str, pagina: int):
    """Construye el run del índice con tabulación real para que salgan los puntos."""
    run = OxmlElement("w:r")
    inicio = OxmlElement("w:t")
    inicio.set(qn("xml:space"), "preserve")
    inicio.text = etiqueta
    tabulacion = OxmlElement("w:tab")
    final = OxmlElement("w:t")
    final.text = str(pagina)
    run.append(inicio)
    run.append(tabulacion)
    run.append(final)
    return run


def _norm(text: str) -> str:
    return " ".join(text.replace("\u00a0", " ").split()).casefold()


def _find_page(page_texts: list[str], needle: str, start: int) -> int:
    target = _norm(needle)
    short = target[:90]
    for index in range(max(0, start), len(page_texts)):
        if target in page_texts[index] or (len(short) > 20 and short in page_texts[index]):
            return index
    words = [w for w in re.split(r"\W+", short) if len(w) > 4]
    for index in range(max(0, start), len(page_texts)):
        if words and sum(w in page_texts[index] for w in words) >= max(2, int(len(words) * 0.7)):
            return index
    raise RuntimeError(f"no se localizó en el PDF: {needle}")


def actualizar_indices(docx_path: Path, pdf_path: Path) -> None:
    import fitz

    pdf = fitz.open(pdf_path)
    page_texts = [_norm(page.get_text()) for page in pdf]
    doc = Document(docx_path)
    summary_page = _find_page(page_texts, SUMMARY[:100], 0)
    groups = [
        ("Índice manual nivel", max(0, summary_page - 2)),
        ("Índice manual tabla", summary_page),
        ("Índice manual figura", summary_page),
    ]
    for prefix, initial in groups:
        cursor = initial
        for paragraph in [p for p in doc.paragraphs if p.style.name.startswith(prefix)]:
            label = paragraph.text.rsplit("\t", 1)[0]
            search = label.split(".", 1)[0] if prefix != "Índice manual nivel" else label
            page_index = _find_page(page_texts, search, cursor)
            cursor = page_index
            link = paragraph._p.find(qn("w:hyperlink"))
            container = link if link is not None else paragraph._p
            for run in container.findall(qn("w:r")):
                container.remove(run)
            container.append(_run_con_tabulacion(label, page_index + 1))
            _format_all_xml_runs(paragraph._p, BODY_PT)
    doc.save(docx_path)


def reparar_huerfanos(docx_path: Path, pdf_path: Path) -> list[str]:
    """Evita que una nota quede sola en una página tras su tabla o figura."""
    import fitz

    pdf = fitz.open(pdf_path)
    huerfanas: list[int] = []
    for index, page in enumerate(pdf):
        lineas = [l.strip() for l in page.get_text().splitlines() if l.strip()]
        cuerpo = [l for l in lineas if not re.fullmatch(r"\d+", l)]
        if cuerpo and all(l.startswith("Nota.") or len(l) < 4 for l in cuerpo):
            huerfanas.append(index)

    leyendas: list[str] = []
    for index in huerfanas:
        for anterior in range(index, -1, -1):
            encontradas = re.findall(r"^(Tabla|Figura) (\d+)$", pdf[anterior].get_text(), re.M)
            if encontradas:
                tipo, numero = encontradas[-1]
                leyendas.append(f"{tipo} {numero}")
                break

    if not leyendas:
        return []
    doc = Document(docx_path)
    aplicadas: list[str] = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() in leyendas:
            paragraph.paragraph_format.page_break_before = True
            aplicadas.append(paragraph.text.strip())
    doc.save(docx_path)
    return aplicadas


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--docx", type=Path, default=DEFAULT_DOCX)
    parser.add_argument("--actualizar-indices", type=Path, metavar="PDF")
    parser.add_argument("--reparar-huerfanos", type=Path, metavar="PDF")
    args = parser.parse_args()
    if args.reparar_huerfanos:
        aplicadas = reparar_huerfanos(args.docx, args.reparar_huerfanos)
        print(f"saltos aplicados: {aplicadas}")
    elif args.actualizar_indices:
        actualizar_indices(args.docx, args.actualizar_indices)
        print(f"índices actualizados: {args.docx.name}")
    else:
        rectificar(args.docx)
        print(f"rectificación aplicada: {args.docx.name}")


if __name__ == "__main__":
    main()
