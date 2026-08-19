#!/usr/bin/env python3
"""Genera copias auditadas del DOCX y XLSM sin sobrescribir los originales."""

from __future__ import annotations

import copy
import os
import shutil
import zipfile
from pathlib import Path
from typing import Iterable
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]
DOCX_ORIGINAL = ROOT / "Tesis jaime fredy horacio avance 16.docx"
DOCX_AUDITADO = ROOT / "Tesis jaime fredy horacio avance 16 - VERSION AUDITADA.docx"
XLSM_ORIGINAL = ROOT / "Estudio Financiero @risk Jaime Horacio Fredy (1).xlsm"
XLSM_AUDITADO = ROOT / "Estudio Financiero @risk Jaime Horacio Fredy (1) - VERSION AUDITADA.xlsm"

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
S_NS = "http://schemas.openxmlformats.org/spreadsheetml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
PKG_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"


def _set_border(container, edge: str, val: str, size: str = "0") -> None:
    borders = container.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        container.append(borders)
    element = borders.find(qn(f"w:{edge}"))
    if element is None:
        element = OxmlElement(f"w:{edge}")
        borders.append(element)
    element.set(qn("w:val"), val)
    element.set(qn("w:sz"), size)
    element.set(qn("w:space"), "0")
    element.set(qn("w:color"), "000000")


def _set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("left", "right", "insideH", "insideV"):
        item = borders.find(qn(f"w:{edge}"))
        if item is None:
            item = OxmlElement(f"w:{edge}")
            borders.append(item)
        item.set(qn("w:val"), "nil")
    for edge in ("top", "bottom"):
        item = borders.find(qn(f"w:{edge}"))
        if item is None:
            item = OxmlElement(f"w:{edge}")
            borders.append(item)
        item.set(qn("w:val"), "single")
        item.set(qn("w:sz"), "8")
        item.set(qn("w:space"), "0")
        item.set(qn("w:color"), "000000")


def _set_cell_margins(cell, twips: int = 45) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge in ("top", "start", "bottom", "end"):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(twips))
        node.set(qn("w:type"), "dxa")


def _apply_apa_table(table, widths: list[float], left_columns: Iterable[int] = (0,)) -> None:
    """Aplica un formato sobrio compatible con las pautas tabulares de APA 7."""
    left_columns = set(left_columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    _set_table_borders(table)

    font_size = 8.0 if len(widths) >= 10 else 9.0 if len(widths) >= 6 else 10.0
    for r_idx, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        if r_idx == 0:
            header = OxmlElement("w:tblHeader")
            header.set(qn("w:val"), "true")
            tr_pr.append(header)
        for c_idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[c_idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = tc_pr.find(qn("w:shd"))
            if shd is not None:
                tc_pr.remove(shd)
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                _set_border(tc_pr, edge, "nil")
            if r_idx == 0:
                _set_border(tc_pr, "bottom", "single", "8")
            _set_cell_margins(cell, 35 if len(widths) >= 10 else 50)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1
                if r_idx == 0:
                    paragraph.alignment = (
                        WD_ALIGN_PARAGRAPH.LEFT if c_idx in left_columns else WD_ALIGN_PARAGRAPH.CENTER
                    )
                else:
                    paragraph.alignment = (
                        WD_ALIGN_PARAGRAPH.LEFT if c_idx in left_columns else WD_ALIGN_PARAGRAPH.RIGHT
                    )
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(font_size)
                    run.bold = r_idx == 0


def _replace_table(doc: Document, old_table, data: list[list[str]], widths: list[float], left_columns=(0,)):
    if not data or any(len(row) != len(data[0]) for row in data):
        raise ValueError("La matriz de reemplazo debe ser rectangular y no vacía")
    if len(widths) != len(data[0]):
        raise ValueError("El número de anchos no coincide con las columnas")
    new_table = doc.add_table(rows=len(data), cols=len(data[0]))
    for r_idx, values in enumerate(data):
        for c_idx, value in enumerate(values):
            new_table.cell(r_idx, c_idx).text = str(value)
    old_table._tbl.addprevious(new_table._tbl)
    old_table._element.getparent().remove(old_table._element)
    _apply_apa_table(new_table, widths, left_columns)
    return new_table


def _next_paragraph(table, doc: Document) -> Paragraph | None:
    node = table._tbl.getnext()
    while node is not None:
        if node.tag == qn("w:p"):
            return Paragraph(node, doc)
        node = node.getnext()
    return None


def _replace_in_paragraph(paragraph: Paragraph, old: str, new: str) -> bool:
    runs = list(paragraph.runs)
    lengths = [len(run.text) for run in runs]
    text = "".join(run.text for run in runs)
    start = text.find(old)
    if start < 0:
        return False
    end = start + len(old)
    pos = 0
    inserted = False
    for run, length in zip(runs, lengths):
        run_start, run_end = pos, pos + length
        pos = run_end
        if run_end <= start or run_start >= end:
            continue
        local_start = max(0, start - run_start)
        local_end = min(length, end - run_start)
        prefix = run.text[:local_start]
        suffix = run.text[local_end:]
        if not inserted:
            run.text = prefix + new + suffix
            inserted = True
        else:
            run.text = prefix + suffix
    return True


def _matrix(table) -> list[list[str]]:
    return [[cell.text.strip() for cell in row.cells] for row in table.rows]


def generar_docx() -> None:
    shutil.copy2(DOCX_ORIGINAL, DOCX_AUDITADO)
    doc = Document(DOCX_AUDITADO)
    original_tables = list(doc.tables)
    if len(original_tables) != 37:
        raise RuntimeError(f"Se esperaban 37 tablas y se encontraron {len(original_tables)}")
    matrices = {i + 1: _matrix(table) for i, table in enumerate(original_tables)}

    # Tabla 7: cálculo ponderado verificable y estructura sin celdas combinadas.
    t7 = [
        ["Criterio", "Peso", "Comayagua: puntaje", "Comayagua: ponderado", "Villanueva: puntaje", "Villanueva: ponderado"],
        ["Costo", "10 %", "3", "0.30", "1", "0.10"],
        ["Logística", "15 %", "2", "0.30", "2", "0.30"],
        ["Proveedores", "15 %", "3", "0.45", "1", "0.15"],
        ["Competencia", "5 %", "1", "0.05", "1", "0.05"],
        ["Disponibilidad de mano de obra", "10 %", "2", "0.20", "2", "0.20"],
        ["Accesibilidad", "10 %", "3", "0.30", "3", "0.30"],
        ["Flujo de personas", "5 %", "2", "0.10", "2", "0.10"],
        ["Acceso a materia prima", "20 %", "3", "0.60", "1", "0.20"],
        ["Impuestos", "5 %", "2", "0.10", "2", "0.10"],
        ["Permisos", "5 %", "2", "0.10", "2", "0.10"],
        ["Total", "100 %", "", "2.50", "", "1.60"],
    ]
    _replace_table(doc, original_tables[6], t7, [1.65, 0.55, 1.05, 1.10, 1.05, 1.10])

    # Tabla 11: conciliación transparente de la diferencia de L 25,000.
    t11 = [
        ["Concepto", "Valor inicial (L)", "Horizonte (años)", "Amortización anual (L)"],
        ["Permisos de funcionamiento", "14,000", "5", "2,800"],
        ["Equipos de protección", "3,500", "5", "700"],
        ["Software de gestión", "8,500", "5", "1,700"],
        ["Inducción y capacitaciones", "25,000", "5", "5,000"],
        ["Provisión no desagregada del modelo financiero*", "25,000", "5", "5,000"],
        ["Total", "76,000", "", "15,200"],
    ]
    t11_new = _replace_table(doc, original_tables[10], t11, [2.85, 1.20, 1.05, 1.40])
    note = _next_paragraph(t11_new, doc)
    if note is None or not note.text.startswith("Nota."):
        raise RuntimeError("No se localizó la nota de la Tabla 11")
    note.text = (
        "Nota. Elaboración propia, 2026. *La provisión de L 25,000 concilia el desglose disponible "
        "de L 51,000 con los L 76,000 usados en el modelo financiero; su naturaleza requiere respaldo documental."
    )

    # Matrices de las tablas financieras, sin títulos internos ni columnas vacías.
    transformed: dict[int, list[list[str]]] = {}
    transformed[15] = [row[:6] for row in matrices[15][1:]]
    transformed[16] = matrices[16][2:]
    transformed[17] = matrices[17][1:]
    transformed[18] = [row[:3] for row in matrices[18][1:]]
    transformed[19] = matrices[19]
    for number in range(20, 27):
        transformed[number] = matrices[number][1:]

    # Tabla 23: no existe cuota cuando la deuda ya está cancelada.
    cuota = next(row for row in transformed[23] if row[0] == "Cuota total")
    cuota[7] = "0"  # Año 6

    # Tablas 25 y 26: hacer visible el ICS que ya descuentan las fórmulas.
    ics_source = next(row[:12] for row in matrices[24] if row[0] == "IMPUESTO MUNICIPAL (ICS)")
    ics_row = ["(−) Impuesto municipal (ICS)"] + ics_source[1:]
    for number in (25, 26):
        target = transformed[number]
        position = next(i for i, row in enumerate(target) if "Utilidad antes de impuestos" in row[0])
        target.insert(position, copy.deepcopy(ics_row))

    transformed[27] = [row[:5] for row in matrices[27][1:4]]
    transformed[28] = [["Indicador", "Valor"]] + [row[:2] for row in matrices[28][1:]]
    next(row for row in transformed[28] if row[0] == "PRI (años)")[1] = "3.67"
    transformed[29] = [row[:5] for row in matrices[29][1:]]
    transformed[30] = [["Indicador", "Valor"]] + [row[:2] for row in matrices[30][1:]]
    next(row for row in transformed[30] if row[0] == "PRI (años)")[1] = "3.66"

    transformed[31] = [
        ["Aspecto", "Criterio aplicado"],
        ["Impuesto sobre la renta", "25 % más aportación solidaria de 5 % sobre el excedente de L 1,000,000."],
        ["Depreciación y amortización", "Línea recta: maquinaria 10 %; vehículo 20 %; gastos diferidos a cinco años."],
        ["Seguridad social", "IHSS patronal 7.2 % con techo mensual de L 11,109.36 e INFOP 1 % sobre la planilla sin techo."],
        ["Prestaciones", "Doce salarios más aguinaldo y decimocuarto mes: 14 meses de salario al año."],
        ["Impuesto sobre ventas", "La leche fluida pasteurizada se trata como exenta; no se carga ISV a las ventas del modelo."],
        ["Pérdidas fiscales", "Compensación de pérdidas fiscales por un máximo de tres años en el modelo."],
        ["Impuesto municipal", "ICS de Comayagua calculado progresivamente sobre ingresos; bienes inmuebles no cuantificados por falta de avalúo."],
        ["Capital de trabajo", "Un mes de costos operativos, financiado 50 % con préstamo y 50 % con recursos propios; recuperación en Año 10."],
    ]

    transformed[32] = [
        ["Parámetro", "Valor actual", "Clasificación", "Mínimo", "Base", "Máximo", "Justificación"],
        ["Tasa de impuesto sobre la renta", "25 %", "Fija", "—", "—", "—", "Parámetro normativo no simulado."],
        ["Aportación solidaria", "5 %", "Fija", "—", "—", "—", "Parámetro normativo no simulado."],
        ["Umbral de aportación solidaria", "L 1,000,000", "Fija", "—", "—", "—", "Umbral monetario no simulado."],
        ["Tasa de interés del préstamo", "14 %", "Fija", "—", "—", "—", "Tasa contractual asumida como fija."],
        ["Financiamiento con préstamo", "50 %", "Fija", "—", "—", "—", "Decisión de estructura financiera."],
        ["Inflación anual", "4.5 %", "Variable", "3.825 %", "4.5 %", "5.4 %", "Entrada Normal configurada en @RISK."],
        ["Tasa de descuento anual", "15 %", "Fija", "—", "—", "—", "Supuesto de decisión no simulado."],
        ["Costo de leche cruda (L/litro)", "13", "Variable", "11.70", "13", "14.95", "Entrada PERT configurada en @RISK."],
        ["Merma de proceso", "3 %", "Variable", "2.4 %", "3 %", "3.9 %", "Entrada PERT configurada en @RISK."],
        ["Costo de empaque (L/litro)", "0.8667", "Variable", "0.78", "0.8667", "0.9533", "Entrada PERT configurada en @RISK."],
        ["Combustible de reparto (L/litro)", "0.8667", "Variable", "0.80", "0.8667", "0.95", "Entrada PERT configurada en @RISK."],
        ["Meses de salario al año", "14", "Fija", "—", "—", "—", "Prestaciones incluidas como supuesto fijo."],
        ["Aporte patronal IHSS", "7.2 %", "Fija", "—", "—", "—", "Parámetro normativo no simulado."],
        ["INFOP patronal", "1 %", "Fija", "—", "—", "—", "Parámetro normativo no simulado."],
        ["Techo mensual IHSS", "L 11,109.36", "Fija", "—", "—", "—", "Techo de cotización no simulado."],
        ["Costos fijos operativos (mes)", "L 28,000", "Variable", "25,200", "28,000", "32,200", "Entrada PERT configurada en @RISK."],
        ["Gastos administrativos (mes)", "L 21,600", "Fija", "—", "—", "—", "No se simula en el libro."],
        ["Consumo por cliente (L/mes)", "150", "Variable", "127.5", "150", "180", "Entrada PERT configurada en @RISK."],
        ["Precio de venta, Año 1 (L/litro)", "28", "Variable", "26", "28", "30", "Entrada PERT configurada en @RISK."],
        ["Clientes activos, Año 1", "60", "Variable", "45", "60", "75", "Entrada PERT configurada en @RISK."],
        ["Clientes activos, Año 5", "280", "Variable", "224", "280", "336", "Entrada PERT configurada en @RISK."],
        ["Clientes activos, Año 10", "690", "Variable", "552", "690", "828", "Entrada PERT configurada en @RISK."],
        ["Meses de capital de trabajo", "1", "Fija", "—", "—", "—", "Política interna no simulada."],
    ]

    transformed[33] = [
        ["Parámetro", "Base", "Escenario mínimo", "Escenario máximo", "Distribución configurada", "Unidad"],
        ["Inflación anual", "4.5 %", "3.825 %", "5.4 %", "Normal (media 4.5 %; DE 0.6 %)", "%"],
        ["Costo de leche cruda", "13", "11.70", "14.95", "PERT", "L/litro"],
        ["Merma de proceso", "3 %", "2.4 %", "3.9 %", "PERT", "%"],
        ["Costo de empaque", "0.8667", "0.78", "0.9533", "PERT", "L/litro"],
        ["Combustible de reparto", "0.8667", "0.80", "0.95", "PERT", "L/litro"],
        ["Costos fijos operativos", "28,000", "25,200", "32,200", "PERT", "L/mes"],
        ["Consumo por cliente", "150", "127.5", "180", "PERT", "L/mes"],
        ["Precio de venta, Año 1", "28", "26", "30", "PERT", "L/litro"],
        ["Clientes activos, Año 1", "60", "45", "75", "PERT", "clientes"],
        ["Clientes activos, Año 5", "280", "224", "336", "PERT", "clientes"],
        ["Clientes activos, Año 10", "690", "552", "828", "PERT", "clientes"],
    ]

    widths_by_table = {
        15: [2.20, 0.75, 0.90, 0.90, 0.75, 0.95],
        16: [2.05] + [0.40] * 11,
        17: [2.05] + [0.40] * 11,
        18: [3.30, 1.35, 1.85],
        19: [2.05] + [0.40] * 11,
        20: [2.05] + [0.40] * 11,
        21: [2.05] + [0.40] * 11,
        22: [2.05] + [0.40] * 11,
        23: [2.05] + [0.40] * 11,
        24: [2.05] + [0.40] * 11,
        25: [2.05] + [0.40] * 11,
        26: [2.05] + [0.40] * 11,
        27: [1.90, 1.25, 1.15, 1.00, 1.20],
        28: [3.80, 2.70],
        29: [1.90, 1.25, 1.15, 1.00, 1.20],
        30: [3.80, 2.70],
        31: [1.55, 4.95],
        32: [1.60, 0.75, 0.75, 0.65, 0.65, 0.65, 1.45],
        33: [1.55, 0.70, 0.85, 0.85, 1.55, 1.00],
    }
    left_columns_by_table = {31: (0, 1), 32: (0, 6), 33: (0, 4, 5)}
    for number in range(15, 34):
        _replace_table(
            doc,
            original_tables[number - 1],
            transformed[number],
            widths_by_table[number],
            left_columns_by_table.get(number, (0,)),
        )

    # Notas de trazabilidad para las tablas sincronizadas con el libro.
    for number in (32, 33):
        # La referencia original sigue siendo válida aunque la tabla se haya sustituido.
        # Se localiza por la posición de la tabla actual en el cuerpo.
        current_table = list(doc.tables)[number - 1]
        current_note = _next_paragraph(current_table, doc)
        if current_note is None or not current_note.text.startswith("Nota."):
            raise RuntimeError(f"No se localizó la nota de la Tabla {number}")
        current_note.text = "Nota. Elaboración propia con base en la configuración de la hoja Simulacion @RISK del modelo financiero, 2026."

    edits = [
        (1647,
         "El Período de Recuperación de la Inversión es de 4.67 años, lo que significa que la inversión inicial es recuperada aproximadamente durante el quinto año de operación.",
         "El Período de Recuperación de la Inversión es de 3.67 años, lo que significa que la inversión inicial se recupera durante el cuarto año de operación."),
        (1655,
         "Los resultados obtenidos mediante los indicadores financieros demuestran que el proyecto es económicamente viable. El Valor Actual Neto positivo de L 15,370,631, la Tasa Interna de Retorno del 55.95 %, el período de recuperación de 4.67 años y la relación Beneficio/Costo superior a uno evidencian que la inversión genera rentabilidad suficiente para compensar el riesgo asumido por los inversionistas.",
         "Los resultados obtenidos mediante los indicadores financieros demuestran que el proyecto es económicamente viable. El Valor Actual Neto positivo de L 15,370,631, la Tasa Interna de Retorno del 55.95 %, el período de recuperación de 3.67 años y la relación Beneficio/Costo superior a uno evidencian que la inversión genera rentabilidad suficiente para compensar el riesgo asumido por los inversionistas."),
        (1683,
         "El período de recuperación de la inversión es de 4.66 años, lo que significa que los recursos aportados por los accionistas son recuperados antes del quinto año de operación. Este resultado es favorable considerando que el horizonte de evaluación es de diez años.",
         "El período de recuperación de la inversión es de 3.66 años, lo que significa que los recursos aportados por los accionistas se recuperan durante el cuarto año de operación. Este resultado es favorable considerando que el horizonte de evaluación es de diez años."),
        (1716,
         "La evaluación financiera sin financiamiento presentó un VAN de L 15,370,631, una TIR de 55.95 %, un PRI de 4.67 años y una Relación Beneficio/Costo de 1.24. Por su parte, la evaluación con financiamiento registró un VAN de L 15,459,686, una TIR de 63.62 %, un PRI de 4.66 años y una Relación Beneficio/Costo de 1.24.",
         "La evaluación financiera sin financiamiento presentó un VAN de L 15,370,631, una TIR de 55.95 %, un PRI de 3.67 años y una Relación Beneficio/Costo de 1.24. Por su parte, la evaluación con financiamiento registró un VAN de L 15,459,686, una TIR de 63.62 %, un PRI de 3.66 años y una Relación Beneficio/Costo de 1.24."),
        (1729,
         "Entre las variables clasificadas como sensibles destacan la inflación anual, el costo de la leche cruda, la merma de proceso, el costo de empaque, el costo de combustible de reparto, los costos fijos operativos y el consumo promedio por cliente. Estas variables dependen de condiciones de mercado, comportamiento de los consumidores, eficiencia operativa y factores macroeconómicos que escapan al control directo de la empresa.",
         "Entre las variables clasificadas como sensibles destacan la inflación anual, el costo de la leche cruda, la merma de proceso, el costo de empaque, el costo de combustible de reparto, los costos fijos operativos, el consumo promedio por cliente, el precio de venta del primer año y los clientes activos de los años 1, 5 y 10. Estas variables dependen de condiciones de mercado, comportamiento de los consumidores, eficiencia operativa y factores macroeconómicos que escapan al control directo de la empresa."),
        (1749,
         "Entre las variables identificadas se encuentran la inflación anual, el costo de la leche cruda, la merma de proceso, el costo de empaque, el costo de combustible de reparto, los costos fijos operativos y el consumo promedio por cliente. Todas ellas poseen una relación directa con los costos de operación, los ingresos o la rentabilidad del proyecto.",
         "El modelo configura once entradas probabilísticas: inflación anual, costo de la leche cruda, merma de proceso, costo de empaque, combustible de reparto, costos fijos operativos, consumo promedio por cliente, precio de venta del primer año y clientes activos de los años 1, 5 y 10. Todas poseen una relación directa con los costos, los ingresos o la rentabilidad del proyecto."),
        (1750,
         "Para la mayoría de los parámetros se recomienda la utilización de distribuciones tipo PERT, debido a que permiten trabajar con un valor mínimo, un valor más probable y un valor máximo, proporcionando una representación adecuada cuando no se dispone de grandes volúmenes de información histórica. En el caso de la inflación anual también puede utilizarse una distribución normal o PERT, considerando los rangos macroeconómicos estimados.",
         "El libro utiliza distribuciones PERT para diez parámetros, definidas por un valor mínimo, uno más probable y uno máximo. Para la inflación anual utiliza una distribución Normal con media de 4.5 % y desviación estándar de 0.6 %; los valores mínimo y máximo de la tabla representan escenarios de referencia y no límites de truncamiento de esa distribución."),
        (1753,
         "La inflación fue modelada mediante una distribución Normal o PERT con valores comprendidos entre 3.8 % y 5.4 %, tomando como valor central el 4.5 %. Esta variable afecta tanto los costos operativos como la evolución futura de los precios de venta.",
         "La inflación se modela mediante una distribución Normal con media de 4.5 % y desviación estándar de 0.6 %. Los escenarios de referencia del libro son 3.825 % y 5.4 %. Esta variable afecta tanto los costos operativos como la evolución futura de los precios de venta."),
        (1755,
         "Se utilizó una distribución PERT con un mínimo de L 12, un valor más probable de L 13 y un máximo de L 15 por litro. Esta variable constituye el principal componente del costo de producción y representa uno de los factores de mayor impacto sobre la rentabilidad del proyecto.",
         "Se utiliza una distribución PERT con un mínimo de L 11.70, un valor más probable de L 13 y un máximo de L 14.95 por litro. Esta variable constituye el principal componente del costo de producción y representa uno de los factores de mayor impacto sobre la rentabilidad del proyecto."),
        (1757,
         "La merma fue modelada mediante una distribución PERT entre 2 % y 4 %, considerando un valor esperado de 3 %. Su comportamiento depende directamente de la eficiencia operativa y de las condiciones reales de producción.",
         "La merma se modela mediante una distribución PERT entre 2.4 % y 3.9 %, con un valor más probable de 3 %. Su comportamiento depende directamente de la eficiencia operativa y de las condiciones reales de producción."),
        (1759,
         "El costo de empaque se modeló mediante una distribución PERT con un rango entre L 0.78 y L 0.95 por litro, debido a la variabilidad observada en los precios de los materiales utilizados para el envasado",
         "El costo de empaque se modela mediante una distribución PERT con un mínimo de L 0.78, un valor más probable de L 0.8667 y un máximo de L 0.9533 por litro, de acuerdo con la configuración del libro."),
        (1768,
         "El consumo promedio por cliente se modeló mediante una distribución PERT con valores comprendidos entre 128 y 180 litros mensuales, manteniendo un valor más probable de 150 litros. Esta variable tiene un efecto directo sobre las ventas proyectadas y los ingresos generados por la empresa.",
         "El consumo promedio por cliente se modela mediante una distribución PERT entre 127.5 y 180 litros mensuales, con un valor más probable de 150 litros. El libro también modela mediante PERT el precio de venta del primer año y los clientes activos de los años 1, 5 y 10, variables que inciden directamente en las ventas proyectadas."),
    ]
    failures = []
    for index, old, new in edits:
        if not _replace_in_paragraph(doc.paragraphs[index], old, new):
            failures.append(index)
    if failures:
        raise RuntimeError(f"No se aplicaron las ediciones de párrafo: {failures}")

    doc.core_properties.title = "Estudio de factibilidad — versión auditada"
    doc.core_properties.subject = "Correcciones demostrables y tablas APA 7"
    doc.core_properties.comments = (
        "Copia auditada. Los originales se conservaron sin cambios; véase INFORME_AUDITORIA.md."
    )
    doc.save(DOCX_AUDITADO)


def _sheet_paths(members: dict[str, bytes]) -> dict[str, str]:
    workbook = ET.fromstring(members["xl/workbook.xml"])
    rels = ET.fromstring(members["xl/_rels/workbook.xml.rels"])
    targets = {
        rel.attrib["Id"]: rel.attrib["Target"]
        for rel in rels.findall(f"{{{PKG_REL_NS}}}Relationship")
    }
    result = {}
    for sheet in workbook.find(f"{{{S_NS}}}sheets"):
        name = sheet.attrib["name"]
        target = targets[sheet.attrib[f"{{{R_NS}}}id"]].lstrip("/")
        if not target.startswith("xl/"):
            target = "xl/" + target
        result[name] = os.path.normpath(target).replace("\\", "/")
    return result


def _cell(root: ET.Element, reference: str) -> ET.Element:
    for cell in root.findall(f".//{{{S_NS}}}c"):
        if cell.attrib.get("r") == reference:
            return cell
    raise KeyError(f"No se encontró la celda {reference}")


def _set_formula(root: ET.Element, reference: str, formula: str, cached: str | None = None) -> None:
    cell = _cell(root, reference)
    f = cell.find(f"{{{S_NS}}}f")
    if f is None:
        f = ET.Element(f"{{{S_NS}}}f")
        cell.insert(0, f)
    f.text = formula
    if cached is not None:
        cell.attrib.pop("t", None)
        value = cell.find(f"{{{S_NS}}}v")
        if value is None:
            value = ET.SubElement(cell, f"{{{S_NS}}}v")
        value.text = cached


def _col_name(number: int) -> str:
    name = ""
    while number:
        number, remainder = divmod(number - 1, 26)
        name = chr(65 + remainder) + name
    return name


def generar_xlsm() -> None:
    shutil.copy2(XLSM_ORIGINAL, XLSM_AUDITADO)
    with zipfile.ZipFile(XLSM_AUDITADO, "r") as source:
        infos = source.infolist()
        members = {info.filename: source.read(info.filename) for info in infos}
        archive_comment = source.comment

    paths = _sheet_paths(members)
    financial_path = paths["Estudio Financiero Det"]
    risk_path = paths["Modelo Riesgo Operativo"]

    financial = ET.fromstring(members[financial_path])
    _set_formula(
        financial,
        "H93",
        'IF(ROUND(H90,2)>0,D35*$B$9/(1-(1+$B$9)^-5),0)',
        "0",
    )
    _set_formula(
        financial,
        "B153",
        'IF(B122>=0,0,IF(C122>=0,-B122/C121,IF(L122<0,">10",COUNTIF(C122:K122,"<0")+(-INDEX(C122:K122,1,COUNTIF(C122:K122,"<0")))/INDEX(C121:L121,1,COUNTIF(C122:K122,"<0")+1))))',
        "3.6712363740691614",
    )
    _set_formula(
        financial,
        "B166",
        'IF(B143>=0,0,IF(C143>=0,-B143/C142,IF(L143<0,">10",COUNTIF(C143:K143,"<0")+(-INDEX(C143:K143,1,COUNTIF(C143:K143,"<0")))/INDEX(C142:L142,1,COUNTIF(C143:K143,"<0")+1))))',
        "3.6592994159444396",
    )
    ET.register_namespace("", S_NS)
    ET.register_namespace("r", R_NS)
    members[financial_path] = ET.tostring(financial, encoding="utf-8", xml_declaration=True)

    risk = ET.fromstring(members[risk_path])
    changed = 0
    for row in range(17, 28):
        probability_row = row - 14
        for column in range(14, 24):  # N:W
            col = _col_name(column)
            _set_formula(risk, f"{col}{row}", f"+_xll.RiskBinomial(1,{col}{probability_row})")
            changed += 1
    if changed != 110:
        raise RuntimeError(f"Se esperaban 110 fórmulas de frecuencia y se cambiaron {changed}")
    members[risk_path] = ET.tostring(risk, encoding="utf-8", xml_declaration=True)

    workbook = ET.fromstring(members["xl/workbook.xml"])
    calc_pr = workbook.find(f"{{{S_NS}}}calcPr")
    if calc_pr is None:
        calc_pr = ET.SubElement(workbook, f"{{{S_NS}}}calcPr")
    calc_pr.set("calcMode", "auto")
    calc_pr.set("fullCalcOnLoad", "1")
    calc_pr.set("forceFullCalc", "1")
    members["xl/workbook.xml"] = ET.tostring(workbook, encoding="utf-8", xml_declaration=True)

    temporary = XLSM_AUDITADO.with_suffix(".tmp.xlsm")
    with zipfile.ZipFile(temporary, "w", allowZip64=True) as target:
        target.comment = archive_comment
        for info in infos:
            target.writestr(info, members[info.filename])
    os.replace(temporary, XLSM_AUDITADO)


def main() -> None:
    generar_docx()
    generar_xlsm()
    print(f"DOCX generado: {DOCX_AUDITADO.name}")
    print(f"XLSM generado: {XLSM_AUDITADO.name}")


if __name__ == "__main__":
    main()
